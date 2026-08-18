"""Prueba del contrato HTTP mínimo del normalizador laboral."""

from __future__ import annotations

import time
from dataclasses import replace
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from agente.api import normalizador, servidor
from agente.normalizador import ejecuciones
from agente.normalizador.ejecuciones import GestorEjecuciones
from agente.normalizador.modelos import (
    Hallazgo,
    ProgresoLimpiezaLLM,
    ResultadoLimpiezaSilabos,
    ResultadoValidacionSilabos,
    UltimoChunkLimpiezaLLM,
)


def _fuente_xlsx() -> bytes:
    """Construye una fuente mínima válida para el contrato HTTP."""

    libro = Workbook()
    convenios = libro.active
    assert convenios is not None
    convenios.title = "Convenios 2030"
    convenios.append(["RUC", "Empresa", "Facultad", "Cód_carrera", "Carrera", "Ciclo_convenio"])
    convenios.append(["20100000000", "Empresa", "Facultad", "01", "Administración", "6"])

    informes = libro.create_sheet("Informes 2030")
    informes.append(
        [
            "Año",
            "Ciclo",
            "Facultad",
            "Cód_carrera",
            "Carrera",
            "Desempeño General",
            "COMPET Adapta bilidad",
            "COMPET Capac. aprender",
            "COMPET Capac. análisis",
            "COMPET Nivel conoci.",
            "COMPET Aplic. conoci.",
            "COMPET Dinamis energía",
            "COMPET Iniciativa autono.",
            "COMPET Creatividad",
            "COMPET Toleran. presión",
            "COMPET Resoluci. problema",
            "COMPET Preocupa orden",
            "COMPET Visión futuro",
            "COMPET Orienta. cliente",
            "COMPET Relacion Interpers.",
            "COMPET Trabajo equipo",
            "COMPET Otros 1",
            "COMPET Califica 1",
            "COMPET Otros 2",
            "COMPET Califica 2",
            "VALORES Etica",
            "VALORES Responsa",
            "VALORES Lealtad",
            "VALORES Adhesion normas",
            "VALORES Puntualidad",
            "VALORES Orienta servicio",
            "ACTITUD Entusiasmo",
            "ACTITUD Responsa. Social",
            "ACTITUD Persistencia",
            "ACTITUD Flexibilidad",
            "Sugerencias y Recomendaciones",
            "RUC",
            "Razon social",
            "Fecha Inicio Aprobada",
            "Fecha Fin Aprobada",
            "Funciones Iniciales",
            "Funciones Finales",
            "Estado",
            "Fch.Prest.Inf.Fin",
            "Fe.Hr.Crea Inf Ini",
            "ciclo creado inf inicial",
            "Ciclo aprobado",
            "TOP 1000 2024.id",
        ]
    )
    informes.append(
        [
            "2030",
            "6",
            "Facultad",
            "01",
            "Administración",
            "Muy Satisfecho",
            "Excelente",
        ]
        + [""] * 29
        + [
            "20100000000",
            "Empresa Uno",
            "",
            "",
            "Analizar datos usando SQL",
            "Analizar datos usando SQL",
            "",
            "2030-01-31",
            "",
            "2030-0",
            "2030-0",
            "",
        ]
    )

    publicaciones = libro.create_sheet("Publicaciones 2030")
    publicaciones.append(
        [
            "Año de la vacante",
            "RUC",
            "Razón Social",
            "Identificación",
            "Tipo de puesto",
            "Fecha de publicación",
            "Fecha de finalización",
            "Creado_Empleo",
            "Área",
            "Área específica",
            "Cargo",
            "Cargo específico",
            "Posición a publicar",
            "Carrera resumen",
            "Funciones",
        ]
    )
    publicaciones.append(
        [
            "2030",
            "20100000000",
            "Empresa Uno",
            "dato",
            "Empleos",
            "",
            "",
            "",
            "Sistemas",
            "",
            "Analista",
            "",
            "Analista de datos",
            "Sistemas",
            "Analizar datos usando SQL",
        ]
    )

    buffer = BytesIO()
    libro.save(buffer)
    return buffer.getvalue()


def _fuente_docx() -> bytes:
    """Construye un sílabo mínimo para verificar el contrato multipart."""

    documento = Document()
    metadata = documento.add_table(rows=1, cols=2)
    metadata.cell(0, 0).text = "Curso"
    metadata.cell(0, 1).text = "Diseño de bases de datos"
    sumilla = documento.add_table(rows=2, cols=1)
    sumilla.cell(0, 0).text = "Sumilla"
    sumilla.cell(1, 0).text = "Modelamiento de bases de datos relacionales."
    competencia = documento.add_table(rows=2, cols=3)
    competencia.cell(0, 0).text = "Competencias genéricas"
    competencia.cell(0, 1).text = "Descripción"
    competencia.cell(0, 2).text = "Código"
    competencia.cell(1, 0).text = "Diseño de bases de datos"
    competencia.cell(1, 1).text = "Diseñar estructuras de datos relacionales."
    competencia.cell(1, 2).text = "G1"
    logro = documento.add_table(rows=2, cols=3)
    logro.cell(0, 0).text = "Logro de aprendizaje general"
    logro.cell(0, 1).text = "Descripción"
    logro.cell(0, 2).text = "Competencias"
    logro.cell(1, 0).text = "L1"
    logro.cell(1, 1).text = "Modelar bases de datos relacionales"
    logro.cell(1, 2).text = "G1"
    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def test_inicia_y_consulta_ejecucion(monkeypatch, tmp_path: Path) -> None:
    """El frontend recibe un ID y puede consultar el resultado del trabajo."""

    gestor = GestorEjecuciones(tmp_path)
    monkeypatch.setattr(normalizador, "gestor_ejecuciones", gestor)
    cliente = TestClient(servidor.app)

    tipo_xlsx = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    respuesta = cliente.post(
        "/normalizador/empleabilidad",
        files={"archivo": ("fuente.xlsx", _fuente_xlsx(), tipo_xlsx)},
    )

    assert respuesta.status_code == 202
    id_ejecucion = respuesta.json()["id_ejecucion"]
    estado = respuesta.json()["estado"]
    for _ in range(100):
        estado = cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}").json()["estado"]
        if estado in {
            "normalizado",
            "normalizado_con_advertencias",
            "no_publicado",
            "rechazado",
            "error",
        }:
            break
        time.sleep(0.01)

    assert estado == "normalizado"
    ejecucion = cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}").json()
    assert ejecucion["normalizacion"]["publicable"] is True
    assert ejecucion["normalizacion"]["registros_procesados"] == {
        "publicaciones": 1,
        "informes": 1,
    }
    assert any(
        output["archivo"].endswith("/requerimiento_laboral.csv") for output in ejecucion["outputs"]
    )
    descarga = cliente.get(
        f"/normalizador/ejecuciones/{id_ejecucion}/outputs/salidas/requerimiento_laboral.csv"
    )
    assert descarga.status_code == 200
    assert "attachment" in descarga.headers["content-disposition"]
    assert descarga.content
    assert (
        cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}/outputs/manifest.json").status_code
        == 404
    )
    errores = cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}/errores")
    assert errores.status_code == 200
    assert errores.json()["hallazgos"] == []


def test_inicia_y_consulta_ejecucion_de_silabos(monkeypatch, tmp_path: Path) -> None:
    """La fuente curricular produce los cuatro CSV del contrato."""

    gestor = GestorEjecuciones(tmp_path)
    monkeypatch.setattr(normalizador, "gestor_ejecuciones", gestor)
    cliente = TestClient(servidor.app)

    respuesta = cliente.post(
        "/normalizador/silabos",
        files={
            "archivo": (
                "silabo.docx",
                _fuente_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"carrera": "Ingeniería de Sistemas", "periodo": "2030-1"},
    )

    assert respuesta.status_code == 202
    id_ejecucion = respuesta.json()["id_ejecucion"]
    estado = respuesta.json()["estado"]
    for _ in range(100):
        estado = cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}").json()["estado"]
        if estado in {
            "limpiado",
            "limpiado_con_advertencias",
            "no_publicado",
            "rechazado",
            "error",
        }:
            break
        time.sleep(0.01)

    assert estado in {"limpiado", "limpiado_con_advertencias"}
    ejecucion = cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}").json()
    assert ejecucion["validacion_silabos"]["valida"] is True
    assert ejecucion["limpieza_silabos"]["registros"] == 1
    assert {
        "salidas/catalogo_competencias.csv",
        "salidas/catalogo_habilidades.csv",
        "salidas/catalogo_herramientas.csv",
        "salidas/cobertura_curricular.csv",
    } <= {output["archivo"] for output in ejecucion["outputs"]}
    descarga = cliente.get(
        f"/normalizador/ejecuciones/{id_ejecucion}/outputs/salidas/cobertura_curricular.csv"
    )
    assert descarga.status_code == 200
    assert "attachment" in descarga.headers["content-disposition"]
    assert descarga.content
    cuarentena = cliente.get(f"/normalizador/ejecuciones/{id_ejecucion}/cuarentena")
    assert cuarentena.status_code == 200
    assert cuarentena.json()["total"] == 0
    assert not any(hallazgo["severidad"] == "error" for hallazgo in ejecucion["hallazgos"])


def test_warning_de_ingestion_marca_limpieza_curricular_con_advertencias(
    monkeypatch, tmp_path: Path
) -> None:
    """Un warning de ingreso persiste y determina el estado final aunque limpiar no halle nada."""

    gestor = GestorEjecuciones(tmp_path)
    id_ejecucion, directorio = gestor.crear("silabos", "paquete.zip")
    ejecucion = gestor._obtener_objeto(id_ejecucion)
    warning = Hallazgo(
        codigo="ARCHIVO_NO_CURRICULAR",
        severidad="warning",
        mensaje="El archivo no es DOCX ni PDF y será ignorado.",
    )
    validacion = ResultadoValidacionSilabos(
        archivo="paquete.zip",
        carrera="Ingeniería de Sistemas",
        periodo="2030-1",
        sha256="sha256",
        valida=True,
        archivos=(),
        hallazgos=(warning,),
    )
    limpieza = ResultadoLimpiezaSilabos(
        registros=1,
        outputs=(),
        hallazgos=(),
        publicable=True,
    )

    monkeypatch.setattr(ejecuciones, "validar_silabos", lambda *_: validacion)
    monkeypatch.setattr(ejecuciones, "limpiar_silabos", lambda *_args, **_kwargs: limpieza)
    monkeypatch.setattr(
        ejecuciones,
        "cargar_catalogo",
        lambda: (_ for _ in ()).throw(RuntimeError("catálogo no necesario para esta prueba")),
    )

    gestor._validar_silabos(
        ejecucion,
        tmp_path / "paquete.zip",
        "Ingeniería de Sistemas",
        "2030-1",
    )

    respuesta = gestor.obtener(id_ejecucion)
    manifest = (directorio / "manifest.json").read_text(encoding="utf-8")

    assert respuesta["estado"] == "limpiado_con_advertencias"
    assert respuesta["hallazgos"] == [warning.a_dict()]
    assert respuesta["validacion_silabos"]["hallazgos"] == [warning.a_dict()]
    assert manifest.count("ARCHIVO_NO_CURRICULAR") == 2


def test_persiste_progreso_llm_en_el_manifest_durante_limpieza(monkeypatch, tmp_path: Path) -> None:
    gestor = GestorEjecuciones(tmp_path)
    id_ejecucion, directorio = gestor.crear("silabos", "paquete.zip")
    ejecucion = gestor._obtener_objeto(id_ejecucion)
    validacion = ResultadoValidacionSilabos(
        archivo="paquete.zip",
        carrera="Marketing",
        periodo="2026-1",
        sha256="sha256",
        valida=True,
        archivos=(),
        hallazgos=(),
    )
    limpieza = ResultadoLimpiezaSilabos(
        registros=1,
        outputs=(),
        hallazgos=(),
        publicable=True,
    )

    def limpiar_con_progreso(*_args, **kwargs):
        actualizar = kwargs["al_actualizar_progreso_llm"]
        assert callable(actualizar)
        inicial = kwargs["progreso_inicial"]
        assert inicial is not None
        assert inicial.fase == "preparando"
        assert inicial.silabos_totales == 0
        assert inicial.eventos[-1].mensaje.startswith("Preparando la extracción")
        actualizar(
            ProgresoLimpiezaLLM(
                fase="completado",
                chunks_completados=2,
                chunks_totales=2,
                logros_procesados=12,
                logros_totales=12,
                silabos_procesados=3,
                silabos_totales=3,
                decisiones_cacheadas=4,
                reintentos=1,
                ultimo_chunk=UltimoChunkLimpiezaLLM("inspector", 4, 1),
                reporte_final="disponible",
            )
        )
        en_polling = gestor.obtener(id_ejecucion)
        assert en_polling["estado"] == "limpiando"
        assert en_polling["progreso_llm"] == {
            "fase": "completado",
            "chunks_completados": 2,
            "chunks_totales": 2,
            "logros_procesados": 12,
            "logros_totales": 12,
            "silabos_detectados": 0,
            "silabos_procesados": 3,
            "silabos_totales": 3,
            "decisiones_cacheadas": 4,
            "reintentos": 1,
            "logros_detectados": 0,
            "mensaje": "",
            "ultimo_chunk": {"fase": "inspector", "logros": 4, "silabos": 1},
            "reporte_final": "disponible",
            "eventos": [],
        }
        return limpieza

    monkeypatch.setenv("NORMALIZADOR_CURRICULAR_LLM", "true")
    monkeypatch.setattr(ejecuciones, "validar_silabos", lambda *_: validacion)
    monkeypatch.setattr(ejecuciones, "limpiar_silabos", limpiar_con_progreso)
    monkeypatch.setattr(
        ejecuciones,
        "cargar_catalogo",
        lambda: (_ for _ in ()).throw(RuntimeError("catálogo no necesario para esta prueba")),
    )

    gestor._validar_silabos(ejecucion, tmp_path / "paquete.zip", "Marketing", "2026-1")

    manifest = (directorio / "manifest.json").read_text(encoding="utf-8")
    assert gestor.obtener(id_ejecucion)["progreso_llm"]["reporte_final"] == "disponible"
    assert '"progreso_llm"' in manifest
    assert '"decisiones_cacheadas": 4' in manifest


def test_publica_evento_de_error_sin_perder_historial(monkeypatch, tmp_path: Path) -> None:
    gestor = GestorEjecuciones(tmp_path)
    id_ejecucion, _directorio = gestor.crear("silabos", "paquete.zip")
    ejecucion = gestor._obtener_objeto(id_ejecucion)
    validacion = ResultadoValidacionSilabos(
        archivo="paquete.zip",
        carrera="Marketing",
        periodo="2026-1",
        sha256="sha256",
        valida=True,
        archivos=(),
        hallazgos=(),
    )

    def limpiar_con_error(*_args, **kwargs):
        actualizar = kwargs["al_actualizar_progreso_llm"]
        progreso = replace(
            kwargs["progreso_inicial"],
            fase="analista",
            chunks_completados=1,
            chunks_totales=2,
            logros_procesados=8,
            logros_totales=16,
            silabos_detectados=76,
            silabos_procesados=1,
        ).con_evento("Chunk 1/2 de Analista LLM completado.")
        actualizar(progreso)
        raise RuntimeError("fallo tardío simulado")

    monkeypatch.setenv("NORMALIZADOR_CURRICULAR_LLM", "true")
    monkeypatch.setattr(ejecuciones, "validar_silabos", lambda *_: validacion)
    monkeypatch.setattr(ejecuciones, "limpiar_silabos", limpiar_con_error)
    monkeypatch.setattr(
        ejecuciones,
        "cargar_catalogo",
        lambda: (_ for _ in ()).throw(RuntimeError("catálogo no necesario para esta prueba")),
    )

    gestor._validar_silabos(ejecucion, tmp_path / "paquete.zip", "Marketing", "2026-1")

    progreso_final = gestor.obtener(id_ejecucion)["progreso_llm"]
    assert isinstance(progreso_final, dict)
    assert progreso_final["fase"] == "error"
    assert any(
        evento["mensaje"].startswith("Chunk 1/2") for evento in progreso_final["eventos"]
    )
    assert progreso_final["eventos"][-1]["mensaje"].startswith(
        "La ejecución terminó con error"
    )
    assert progreso_final["silabos_detectados"] == 76
    assert progreso_final["silabos_procesados"] == 1
