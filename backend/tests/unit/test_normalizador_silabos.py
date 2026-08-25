"""Pruebas del contrato y staging curricular."""

from __future__ import annotations

import csv
import json
import zipfile
from dataclasses import replace
from pathlib import Path

from docx import Document

from agente.normalizador.empleabilidad.catalogo import CatalogoCHH, ConceptoCHH
from agente.normalizador.excepciones import CancelacionSolicitada
from agente.normalizador.modelos import ProgresoLimpiezaLLM
from agente.normalizador.silabos import limpieza
from agente.normalizador.silabos.analista_llm import ResultadoAnalisisCurricular
from agente.normalizador.silabos.entrada import validar_archivo
from agente.normalizador.silabos.limpieza import (
    _competencias_pdf,
    _logros_pdf,
    limpiar_archivo,
)
from agente.normalizador.silabos.salida import (
    _competencias_declaradas_por_texto,
    _resolver_habilidad_canonica,
)


def _crear_docx(ruta: Path) -> None:
    documento = Document()
    metadata = documento.add_table(rows=1, cols=2)
    metadata.cell(0, 0).text = "Curso"
    metadata.cell(0, 1).text = "Diseño de bases de datos"
    sumilla = documento.add_table(rows=2, cols=1)
    sumilla.cell(0, 0).text = "Sumilla"
    sumilla.cell(1, 0).text = "Modelamiento de bases de datos relacionales."
    competencia = documento.add_table(rows=2, cols=3)
    competencia.cell(0, 0).text = "Competencias genéricas"
    competencia.cell(0, 1).text = "Descripción"
    competencia.cell(0, 2).text = "Código"
    competencia.cell(1, 0).text = "Diseño de bases de datos"
    competencia.cell(1, 1).text = "Diseñar estructuras de datos relacionales."
    competencia.cell(1, 2).text = "G1"
    logro = documento.add_table(rows=2, cols=3)
    logro.cell(0, 0).text = "Logro de aprendizaje general"
    logro.cell(0, 1).text = "Descripción"
    logro.cell(0, 2).text = "Competencias"
    logro.cell(1, 0).text = "L1"
    logro.cell(1, 1).text = "Modelar bases de datos relacionales"
    logro.cell(1, 2).text = "G1"
    documento.save(ruta)


def _crear_docx_con_logro_verticalmente_combinado(ruta: Path) -> None:
    """Replica una fila L3 cuyo XML usa ``w:vMerge`` en las tres columnas."""

    documento = Document()
    tabla = documento.add_table(rows=4, cols=3)
    for celda, valor in zip(
        tabla.rows[0].cells,
        ("Logro de aprendizaje general", "Descripción", "Competencias"),
        strict=True,
    ):
        celda.text = valor
    tabla.cell(1, 0).text = "L3"
    tabla.cell(1, 1).text = "Desarrollar el texto completo y verificable del logro L3."
    tabla.cell(1, 2).text = "E3"
    for columna in range(3):
        tabla.cell(1, columna).merge(tabla.cell(3, columna))
    documento.save(ruta)


def _crear_docx_codigo_no_declarado(ruta: Path, incluir_competencia: bool = True) -> None:
    documento = Document()
    metadata = documento.add_table(rows=1, cols=2)
    metadata.cell(0, 0).text = "Curso"
    metadata.cell(0, 1).text = "Introducción a las finanzas"
    sumilla = documento.add_table(rows=2, cols=1)
    sumilla.cell(0, 0).text = "Sumilla"
    sumilla.cell(1, 0).text = "Fundamentos de evaluación de indicadores financieros."
    if incluir_competencia:
        competencia = documento.add_table(rows=2, cols=3)
        competencia.cell(0, 0).text = "Competencias específicas"
        competencia.cell(0, 1).text = "Descripción"
        competencia.cell(0, 2).text = "Código"
        competencia.cell(1, 0).text = "Pensamiento crítico"
        competencia.cell(1, 1).text = "Evaluar argumentos y evidencia para sustentar decisiones."
        competencia.cell(1, 2).text = "G7"
    logro = documento.add_table(rows=2, cols=3)
    logro.cell(0, 0).text = "Logro de aprendizaje general"
    logro.cell(0, 1).text = "Descripción"
    logro.cell(0, 2).text = "Competencias"
    logro.cell(1, 0).text = "L2"
    logro.cell(1, 1).text = (
        "Evaluar indicadores financieros para sustentar decisiones de inversión."
    )
    logro.cell(1, 2).text = "E2"
    documento.save(ruta)


def _crear_docx_codigo_alfabetico_con_herramientas(ruta: Path) -> None:
    documento = Document()
    metadata = documento.add_table(rows=1, cols=2)
    metadata.cell(0, 0).text = "Curso"
    metadata.cell(0, 1).text = "Herramientas para marketing"
    sumilla = documento.add_table(rows=2, cols=1)
    sumilla.cell(0, 0).text = "Sumilla"
    sumilla.cell(1, 0).text = "Gestión estratégica de productos y clientes."
    competencia = documento.add_table(rows=2, cols=3)
    competencia.cell(0, 0).text = "Competencias específicas"
    competencia.cell(0, 1).text = "Descripción"
    competencia.cell(0, 2).text = "Código"
    competencia.cell(1, 0).text = "Gestión estratégica"
    competencia.cell(1, 1).text = "Diseñar estrategias de marketing."
    competencia.cell(1, 2).text = "E1"
    logro = documento.add_table(rows=2, cols=3)
    logro.cell(0, 0).text = "Logro de aprendizaje general"
    logro.cell(0, 1).text = "Descripción"
    logro.cell(0, 2).text = "Competencias"
    logro.cell(1, 0).text = "L1"
    logro.cell(1, 1).text = "Diseñar una estrategia de marketing para un producto."
    logro.cell(1, 2).text = "EE"
    recursos = documento.add_table(rows=2, cols=1)
    recursos.cell(0, 0).text = "Recursos de aprendizaje"
    recursos.cell(1, 0).text = "MS Excel, MS Word y MS Project."
    documento.save(ruta)


def _crear_docx_con_bibliografia_que_parece_herramienta(ruta: Path) -> None:
    _crear_docx_codigo_alfabetico_con_herramientas(ruta)
    documento = Document(ruta)
    recursos = documento.tables[-1]
    recursos.cell(1, 0).text = "Microsoft Excel para analizar datos."
    documento.add_paragraph("Bibliografía")
    documento.add_paragraph(
        "https://ejemplo.edu/index.php; Box-Jenkins; Valor Presente Neto (VPN); "
        "Brijs, Bert; Slack, N."
    )
    documento.save(ruta)


def _catalogo_financiero() -> CatalogoCHH:
    return CatalogoCHH(
        competencias=(
            ConceptoCHH(
                "COMP_FIN",
                "Evaluación financiera",
                "Evaluar indicadores financieros para sustentar decisiones.",
                "dura",
            ),
        ),
        habilidades=(
            ConceptoCHH(
                "HAB_FIN",
                "Evaluar indicadores financieros para sustentar decisiones de inversión",
                "Evaluar indicadores financieros para sustentar decisiones de inversión.",
            ),
        ),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )


def _catalogo_con_habilidad_bases_de_datos() -> CatalogoCHH:
    return CatalogoCHH(
        competencias=(),
        habilidades=(
            ConceptoCHH(
                "HAB_DB",
                "Modelar bases de datos relacionales",
                "Modelar bases de datos relacionales.",
            ),
        ),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )


def test_valida_y_limpia_docx_con_carrera_y_periodo(tmp_path: Path) -> None:
    fuente = tmp_path / "Ciclo_03" / "DISENO_DE_BASES_DE_DATOS.docx"
    fuente.parent.mkdir()
    _crear_docx(fuente)
    validacion = validar_archivo(fuente, "Ingeniería de Sistemas", "2030-1")

    assert validacion.valida is True
    assert validacion.carrera == "INGENIERIA_DE_SISTEMAS"
    assert validacion.archivos[0].formato == "docx"

    ejecucion = tmp_path / "ejecucion"
    resultado = limpiar_archivo(
        fuente,
        ejecucion,
        validacion,
        _catalogo_con_habilidad_bases_de_datos(),
    )
    assert resultado.registros == 1
    assert resultado.publicable is True
    assert resultado.relaciones == 1
    registro = json.loads((ejecucion / "limpios" / "silabos.jsonl").read_text(encoding="utf-8"))
    assert registro["datos"]["curso"] == "Diseño de bases de datos"
    assert registro["datos"]["logros_especificos"][0]["etiqueta"] == "L1"
    assert {output["archivo"] for output in resultado.outputs} == {
        "salidas/catalogo_competencias.csv",
        "salidas/catalogo_habilidades.csv",
        "salidas/catalogo_herramientas.csv",
        "salidas/cobertura_curricular.csv",
        "salidas/competencias_fuente.jsonl",
        "salidas/habilidades_fuente.jsonl",
        "salidas/herramientas_fuente.jsonl",
        "salidas/cobertura_curricular_fuente.jsonl",
        "salidas/cobertura_curricular_canonica.jsonl",
        "salidas/pendientes_curriculares.jsonl",
        "salidas/release_gate.json",
    }
    schemas = {
        "catalogo_competencias.csv": [
            "id_competencia",
            "nombre_competencia",
            "descripcion_breve_competencia",
            "tipo_competencia",
        ],
        "catalogo_habilidades.csv": [
            "id_habilidad",
            "nombre_habilidad",
            "descripcion_breve",
        ],
        "catalogo_herramientas.csv": [
            "id_herramienta",
            "nombre_herramienta",
            "descripcion_breve_herramienta",
        ],
        "cobertura_curricular.csv": [
            "id_cob_curricular",
            "id_curso",
            "id_silabo",
            "id_competencia",
            "id_habilidad",
            "id_herramienta",
        ],
    }
    for nombre, esperado in schemas.items():
        with (ejecucion / "salidas" / nombre).open(
            encoding="utf-8-sig", newline=""
        ) as archivo:
            assert next(csv.reader(archivo)) == esperado
    assert (ejecucion / "salidas" / "reportes" / "habilidades_fuente.jsonl").is_file()


def test_publica_logros_detectados_durante_la_extraccion(
    monkeypatch, tmp_path: Path
) -> None:
    fuente = tmp_path / "Ciclo_03" / "DISENO_DE_BASES_DE_DATOS.docx"
    fuente.parent.mkdir()
    _crear_docx(fuente)
    validacion = validar_archivo(fuente, "Ingeniería de Sistemas", "2030-1")
    progresos = []
    monkeypatch.setattr(
        limpieza,
        "analizar_registros_curriculares",
        lambda *_args, **_kwargs: ResultadoAnalisisCurricular(
            decisiones={},
            reportes=(),
            modelo_analista="llm-test",
            modelo_inspector="no_ejecutado",
            lotes=1,
        ),
    )

    limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        _catalogo_con_habilidad_bases_de_datos(),
        usar_llm=True,
        al_actualizar_progreso_llm=progresos.append,
    )

    progreso_extraccion = next(progreso for progreso in progresos if progreso.fase == "extrayendo")
    assert progreso_extraccion.logros_detectados == 1
    assert progreso_extraccion.logros_totales == 1
    assert progreso_extraccion.silabos_detectados == 1
    assert progreso_extraccion.silabos_procesados == 0
    assert progreso_extraccion.silabos_totales == 1
    assert progreso_extraccion.eventos[-1].mensaje == (
        "Logros detectados: 1. Sílabos detectados: 1/1."
    )


def test_preserva_el_ultimo_chunk_si_el_analista_falla_tarde(
    monkeypatch, tmp_path: Path
) -> None:
    fuente = tmp_path / "Ciclo_03" / "DISENO_DE_BASES_DE_DATOS.docx"
    fuente.parent.mkdir()
    _crear_docx(fuente)
    validacion = validar_archivo(fuente, "Ingeniería de Sistemas", "2030-1")
    progresos = []

    def analizar_con_fallo(*_args, **kwargs):
        actualizar = kwargs["al_actualizar_progreso"]
        progreso = ProgresoLimpiezaLLM(
            fase="analista",
            chunks_completados=1,
            chunks_totales=2,
            logros_procesados=8,
            logros_totales=16,
            silabos_procesados=1,
            silabos_totales=76,
            decisiones_cacheadas=8,
            reintentos=0,
            silabos_detectados=76,
        ).con_evento("Chunk 1/2 de Analista LLM completado.")
        actualizar(progreso)
        actualizar(
            replace(
                progreso,
                chunks_completados=2,
                logros_procesados=16,
                silabos_procesados=3,
            ).con_evento("Chunk 2/2 de Analista LLM completado.")
        )
        raise RuntimeError("fallo tardío simulado")

    monkeypatch.setattr(limpieza, "analizar_registros_curriculares", analizar_con_fallo)

    limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        _catalogo_con_habilidad_bases_de_datos(),
        usar_llm=True,
        al_actualizar_progreso_llm=progresos.append,
    )

    progreso_final = progresos[-1]
    assert progreso_final.fase == "error"
    assert progreso_final.silabos_detectados == 76
    assert progreso_final.silabos_procesados == 3
    assert any("Chunk 1/2" in evento.mensaje for evento in progreso_final.eventos)
    assert any("Chunk 2/2" in evento.mensaje for evento in progreso_final.eventos)
    assert progreso_final.eventos[-1].mensaje.startswith("El análisis LLM no estuvo disponible")


def test_cancelacion_conserva_reportes_auditables_llm(
    monkeypatch, tmp_path: Path
) -> None:
    fuente = tmp_path / "Ciclo_03" / "DISENO_DE_BASES_DE_DATOS.docx"
    fuente.parent.mkdir()
    _crear_docx(fuente)
    validacion = validar_archivo(fuente, "Ingeniería de Sistemas", "2030-1")

    monkeypatch.setattr(
        limpieza,
        "analizar_registros_curriculares",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(CancelacionSolicitada()),
    )

    ejecucion = tmp_path / "ejecucion"
    try:
        limpiar_archivo(
            fuente,
            ejecucion,
            validacion,
            _catalogo_con_habilidad_bases_de_datos(),
            usar_llm=True,
        )
    except CancelacionSolicitada:
        pass
    else:
        raise AssertionError("La cancelación debe propagarse al gestor de ejecuciones.")

    reportes = ejecucion / "salidas" / "reportes"
    assert (reportes / "decisiones_llm.jsonl").is_file()
    assert (reportes / "cuarentena.jsonl").is_file()
    analisis = json.loads((reportes / "analisis_llm.json").read_text(encoding="utf-8"))
    assert analisis["estado"] == "CANCELADO"


def test_extrae_una_sola_fila_logica_para_logro_con_vmerge(tmp_path: Path) -> None:
    fuente = tmp_path / "MARKETING_SOCIAL.docx"
    _crear_docx_con_logro_verticalmente_combinado(fuente)

    registro = limpieza._extraer_docx(fuente, fuente.name, "MARKETING", "2026-1")
    datos = registro["datos"]
    assert isinstance(datos, dict)
    logros = datos["logros_especificos"]
    assert isinstance(logros, list)

    assert [logro["etiqueta"] for logro in logros] == ["L3"]
    assert logros[0] == {
        "etiqueta": "L3",
        "descripcion": "Desarrollar el texto completo y verificable del logro L3.",
        "codigos_competencia": ["E3"],
    }


def test_publica_auditoria_contexto_en_resumen_llm(
    monkeypatch, tmp_path: Path
) -> None:
    fuente = tmp_path / "DISENO_DE_BASES_DE_DATOS.docx"
    _crear_docx(fuente)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")
    auditoria = {
        "version_contexto": "contexto-curricular/v1",
        "version_catalogo": "catalogo-test",
        "hash_contextos": "hash-test",
    }
    analisis = ResultadoAnalisisCurricular(
        decisiones={},
        reportes=(),
        modelo_analista="llm-test",
        modelo_inspector="no_ejecutado",
        lotes=1,
        auditoria_contexto=auditoria,
    )
    monkeypatch.setattr(
        limpieza,
        "analizar_registros_curriculares",
        lambda *_args, **_kwargs: analisis,
    )

    limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        _catalogo_con_habilidad_bases_de_datos(),
        usar_llm=True,
    )

    resumen = json.loads(
        (
            tmp_path / "ejecucion" / "salidas" / "reportes" / "analisis_llm.json"
        ).read_text(encoding="utf-8")
    )
    assert resumen["auditoria_contexto"] == auditoria


def test_rechaza_ruta_insegura_en_zip(tmp_path: Path) -> None:
    fuente = tmp_path / "curriculo.zip"
    with zipfile.ZipFile(fuente, "w") as paquete:
        paquete.writestr("../fuera.docx", b"no es un docx")

    validacion = validar_archivo(fuente, "Marketing", "2030-1")

    assert validacion.valida is False
    assert any(hallazgo.codigo == "RUTA_ZIP_INSEGURA" for hallazgo in validacion.hallazgos)


def test_ignora_silenciosamente_metadatos_de_macos_en_zip(tmp_path: Path) -> None:
    docx = tmp_path / "curso.docx"
    _crear_docx(docx)
    fuente = tmp_path / "curriculo.zip"
    with zipfile.ZipFile(fuente, "w") as paquete:
        paquete.write(docx, "MARKETING/curso.docx")
        paquete.writestr("__MACOSX/MARKETING/._curso.docx", b"metadata")

    validacion = validar_archivo(fuente, "Marketing", "2030-1")

    assert validacion.valida is True
    assert [archivo.nombre for archivo in validacion.archivos] == ["MARKETING/curso.docx"]
    assert validacion.hallazgos == ()


def test_conserva_logro_cuando_codigo_no_aparece_en_tabla(tmp_path: Path) -> None:
    fuente = tmp_path / "INTRODUCCION_A_LAS_FINANZAS.docx"
    _crear_docx_codigo_no_declarado(fuente)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")

    resultado = limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        _catalogo_financiero(),
    )

    assert resultado.publicable is True
    assert resultado.relaciones == 1
    assert not any(hallazgo.severidad == "error" for hallazgo in resultado.hallazgos)
    assert any(
        hallazgo.codigo == "COMPETENCIA_REFERENCIADA_NO_DECLARADA"
        for hallazgo in resultado.hallazgos
    )
    assert any(
        hallazgo.codigo == "LOGRO_CODIGO_INCONSISTENTE"
        for hallazgo in resultado.hallazgos
    )
    with (tmp_path / "ejecucion" / "salidas" / "catalogo_competencias.csv").open(
        encoding="utf-8-sig", newline=""
    ) as archivo:
        filas = list(csv.DictReader(archivo))
    assert [fila["nombre_competencia"] for fila in filas] == ["Pensamiento crítico"]
    assert not any(
        fila["nombre_competencia"].startswith("Competencia referenciada")
        for fila in filas
    )
    competencias_fuente = [
        json.loads(line)
        for line in (
            tmp_path
            / "ejecucion"
            / "salidas"
            / "reportes"
            / "competencias_fuente.jsonl"
        ).read_text(encoding="utf-8").splitlines()
    ]
    assert any(
        fila["codigo_competencia"] == "E2"
        and fila["estado_resolucion"] == "REFERENCIA_NO_DECLARADA"
        for fila in competencias_fuente
    )
    cuarentena = tmp_path / "ejecucion" / "salidas" / "reportes" / "cuarentena.jsonl"
    assert cuarentena.read_text(encoding="utf-8").strip() == ""


def test_conserva_referencia_de_fuente_sin_catalogo_o_declaracion(tmp_path: Path) -> None:
    fuente = tmp_path / "CURSO_FINANCIERO.docx"
    _crear_docx_codigo_no_declarado(fuente, incluir_competencia=False)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")
    catalogo_vacio = CatalogoCHH((), (), (), {}, ("test",), "test")

    resultado = limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        catalogo_vacio,
    )

    assert resultado.publicable is True
    assert resultado.relaciones == 0
    assert any(
        hallazgo.codigo == "CURSO_SIN_COMPETENCIA_DECLARADA"
        for hallazgo in resultado.hallazgos
    )
    with (tmp_path / "ejecucion" / "salidas" / "catalogo_competencias.csv").open(
        encoding="utf-8-sig", newline=""
    ) as archivo:
        filas = list(csv.DictReader(archivo))
    assert filas == []
    habilidades_fuente = (
        tmp_path
        / "ejecucion"
        / "salidas"
        / "reportes"
        / "habilidades_fuente.jsonl"
    ).read_text(encoding="utf-8").splitlines()
    assert len(habilidades_fuente) == 1
    cobertura_fuente = (
        tmp_path
        / "ejecucion"
        / "salidas"
        / "reportes"
        / "cobertura_curricular_fuente.jsonl"
    ).read_text(encoding="utf-8").splitlines()
    assert len(cobertura_fuente) == 1


def test_prioriza_perfil_del_silabo_y_conserva_codigo_alfabetico(tmp_path: Path) -> None:
    fuente = tmp_path / "MATEMATICA_PARA_LA_GESTION.docx"
    _crear_docx_codigo_alfabetico_con_herramientas(fuente)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")
    catalogo = CatalogoCHH(
        competencias=(
            ConceptoCHH(
                "COMP_SISTEMAS",
                "Álgebra lineal aplicada",
                "Resolver matrices y sistemas.",
                "dura",
            ),
        ),
        habilidades=(),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )

    resultado = limpiar_archivo(fuente, tmp_path / "ejecucion", validacion, catalogo)

    assert resultado.publicable is True
    assert resultado.competencias == 1
    registro = json.loads(
        (tmp_path / "ejecucion" / "limpios" / "silabos.jsonl").read_text(
            encoding="utf-8"
        )
    )
    assert registro["datos"]["logros_especificos"][0]["codigos_competencia"] == ["EE"]
    with (tmp_path / "ejecucion" / "salidas" / "catalogo_competencias.csv").open(
        encoding="utf-8-sig", newline=""
    ) as archivo:
        filas = list(csv.DictReader(archivo))
    assert filas[0]["nombre_competencia"] == "Gestión estratégica"
    assert not any(
        fila["nombre_competencia"].startswith("Competencia referenciada")
        for fila in filas
    )
    assert "Álgebra lineal aplicada" not in {fila["nombre_competencia"] for fila in filas}
    fuente = (
        tmp_path
        / "ejecucion"
        / "salidas"
        / "reportes"
        / "competencias_fuente.jsonl"
    ).read_text(encoding="utf-8").splitlines()
    assert any(json.loads(line)["codigo_competencia"] == "E1" for line in fuente)


def test_extrae_herramientas_de_recurso_estructurado_y_acepta_alias_ms(tmp_path: Path) -> None:
    fuente = tmp_path / "HERRAMIENTAS_MARKETING.docx"
    _crear_docx_codigo_alfabetico_con_herramientas(fuente)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")
    catalogo = CatalogoCHH(
        competencias=(),
        habilidades=(),
        herramientas=(
            ConceptoCHH("EXCEL", "Microsoft Excel", "Hoja de cálculo"),
            ConceptoCHH("WORD", "Microsoft Word", "Procesador de texto"),
            ConceptoCHH("PROJECT", "Microsoft Project", "Gestión de proyectos"),
        ),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )

    resultado = limpiar_archivo(fuente, tmp_path / "ejecucion", validacion, catalogo)

    assert resultado.publicable is True
    assert resultado.herramientas == 3
    herramientas_fuente = [
        json.loads(line)
        for line in (
            tmp_path
            / "ejecucion"
            / "salidas"
            / "reportes"
            / "herramientas_fuente.jsonl"
        ).read_text(encoding="utf-8").splitlines()
    ]
    assert {fila["seccion_fuente"] for fila in herramientas_fuente} == {
        "recursos de aprendizaje"
    }
    assert all("MS " in fila["texto_evidencia"] for fila in herramientas_fuente)


def test_no_publica_herramientas_encontradas_solo_en_bibliografia(tmp_path: Path) -> None:
    fuente = tmp_path / "BIBLIOGRAFIA_MARKETING.docx"
    _crear_docx_con_bibliografia_que_parece_herramienta(fuente)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")
    catalogo = CatalogoCHH(
        competencias=(),
        habilidades=(),
        herramientas=(
            ConceptoCHH("EXCEL", "Microsoft Excel", "Hoja de cálculo"),
            ConceptoCHH("PHP", "PHP", "Lenguaje"),
            ConceptoCHH("JENKINS", "Jenkins", "Automatización"),
            ConceptoCHH("VPN", "VPN", "Red privada virtual"),
            ConceptoCHH("BERT", "BERT", "Modelo"),
            ConceptoCHH("SLACK", "Slack", "Colaboración"),
        ),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )

    resultado = limpiar_archivo(fuente, tmp_path / "ejecucion", validacion, catalogo)

    assert resultado.publicable is True
    assert resultado.herramientas == 1
    with (tmp_path / "ejecucion" / "salidas" / "catalogo_herramientas.csv").open(
        encoding="utf-8-sig", newline=""
    ) as archivo:
        filas = list(csv.DictReader(archivo))
    assert [fila["nombre_herramienta"] for fila in filas] == ["Microsoft Excel"]


def test_resuelve_habilidad_por_descripcion_y_rechaza_empate() -> None:
    descripcion = "Analizar comportamiento de consumidores mediante métricas de mercado."
    catalogo = CatalogoCHH(
        competencias=(),
        habilidades=(
            ConceptoCHH(
                "HAB_OK",
                "Gestión de información comercial",
                descripcion,
            ),
        ),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )

    resolucion = _resolver_habilidad_canonica(catalogo, descripcion)

    assert resolucion.concepto is not None
    assert resolucion.concepto.id == "HAB_OK"
    assert resolucion.metodo == "COINCIDENCIA_DESCRIPCION"
    assert resolucion.puntaje == 1.0

    ambigua = CatalogoCHH(
        competencias=(),
        habilidades=(
            ConceptoCHH("HAB_A", "Análisis comercial", descripcion),
            ConceptoCHH("HAB_B", "Investigación de mercado", descripcion),
        ),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="test",
    )
    resolucion_ambigua = _resolver_habilidad_canonica(ambigua, descripcion)
    assert resolucion_ambigua.concepto is None
    assert resolucion_ambigua.metodo == "AMBIGUA_O_INSUFICIENTE"


def test_fallback_competencia_exige_margen_entre_candidatas() -> None:
    declaraciones = [
        {
            "codigo": "E1",
            "nombre": "Gestión estratégica",
            "descripcion": "Analizar datos de mercado para decisiones.",
        },
        {
            "codigo": "E2",
            "nombre": "Pensamiento crítico",
            "descripcion": "Analizar datos de mercado para decisiones.",
        },
    ]

    assert _competencias_declaradas_por_texto(
        declaraciones,
        {"curso": "Marketing", "texto_relevante": ""},
        "Analizar datos de mercado para decisiones.",
    ) == []


def test_parser_pdf_admite_codigos_alfabeticos_sin_confundir_palabras() -> None:
    competencias = _competencias_pdf(
        "Competencias específicas Gestión estratégica Diseñar estrategias de marketing EE"
    )
    logros = _logros_pdf(
        "Logros de aprendizaje específicos L1 Diseñar una estrategia de marketing EE"
    )

    assert competencias[0]["codigo"] == "EE"
    assert logros == [
        {
            "etiqueta": "L1",
            "descripcion": "Diseñar una estrategia de marketing",
            "codigos_competencia": ["EE"],
        }
    ]
