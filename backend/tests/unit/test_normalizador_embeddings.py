"""Pruebas deterministas del primer corte de recuperación semántica."""

from __future__ import annotations

import json

import pytest
from docx import Document

from agente.normalizador.embeddings import (
    FALLBACK_REASON_CANDIDATES_BELOW_THRESHOLD,
    FALLBACK_REASON_CATALOG_EMPTY,
    FALLBACK_REASON_PROVIDER_OR_VECTOR_INVALID,
    FALLBACK_REASON_RETRIEVER_ABSENT,
    CatalogDocument,
    EmbeddingRetriever,
    EmbeddingScope,
    EmbeddingUnavailable,
    InMemoryEmbeddingIndex,
    crear_retriever_curricular_opt_in,
)
from agente.normalizador.empleabilidad.catalogo import CatalogoCHH, ConceptoCHH
from agente.normalizador.silabos import analista_llm
from agente.normalizador.silabos import limpieza as limpieza_silabos
from agente.normalizador.silabos.analista_llm import ResultadoAnalisisCurricular
from agente.normalizador.silabos.contexto_curricular import construir_contexto_por_logro
from agente.normalizador.silabos.entrada import validar_archivo
from agente.normalizador.silabos.limpieza import limpiar_archivo


class _FakeEmbeddingProvider:
    model_name = "fake-embeddings-v1"
    config_identifier = "fake-config"

    def __init__(self) -> None:
        self.queries: list[str] = []

    def embed_query(self, text: str) -> list[float]:
        self.queries.append(text)
        return [0.0, 1.0] if "analizar" in text.lower() else [1.0, 0.0]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [
            [0.0, 1.0] if "analizar" in text.lower() else [1.0, 0.0]
            for text in texts
        ]


class _FailingEmbeddingProvider(_FakeEmbeddingProvider):
    def embed_query(self, text: str) -> list[float]:
        raise RuntimeError("provider unavailable")


class _VectorsProvider(_FakeEmbeddingProvider):
    def __init__(self, vectors: list[list[float]]) -> None:
        super().__init__()
        self.vectors = vectors

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.vectors[: len(texts)]


def _documents() -> tuple[CatalogDocument, ...]:
    return (
        CatalogDocument(
            id="HAB_ANALIZAR",
            text="Analizar datos de marketing",
            source_kind="career_curriculum",
            career="MARKETING",
            period="2026-1",
            type="habilidad",
            catalog_version="career-v1",
        ),
        CatalogDocument(
            id="HAB_DISEÑAR",
            text="Diseñar campañas de marketing",
            source_kind="career_curriculum",
            career="MARKETING",
            period="2026-1",
            type="habilidad",
            catalog_version="career-v1",
        ),
        CatalogDocument(
            id="LAB_ANALIZAR",
            text="Analizar indicadores laborales",
            source_kind="labor",
            career=None,
            period=None,
            type="habilidad",
            catalog_version="labor-v1",
        ),
        CatalogDocument(
            id="OTHER_CAREER",
            text="Analizar datos de ingeniería",
            source_kind="career_curriculum",
            career="INGENIERIA",
            period="2026-1",
            type="habilidad",
            catalog_version="career-v2",
        ),
    )


def test_dos_logros_consultan_y_recuperan_resultados_distintos() -> None:
    provider = _FakeEmbeddingProvider()
    retriever = EmbeddingRetriever(provider, InMemoryEmbeddingIndex(_documents()))
    scope = EmbeddingScope(career="MARKETING", period="2026-1")

    primero = retriever.retrieve("Diseñar campañas", scope=scope, limits={"habilidad": 2})
    segundo = retriever.retrieve("Analizar datos", scope=scope, limits={"habilidad": 2})

    assert provider.queries == ["Diseñar campañas", "Analizar datos"]
    assert primero["habilidad"][0].document.id == "HAB_DISEÑAR"
    assert segundo["habilidad"][0].document.id == "HAB_ANALIZAR"


def test_scope_curricular_excluye_labor_y_otras_carreras_salvo_solicitud_explicita() -> None:
    retriever = EmbeddingRetriever(_FakeEmbeddingProvider(), InMemoryEmbeddingIndex(_documents()))
    curricular = retriever.retrieve(
        "Analizar datos",
        scope=EmbeddingScope(career="MARKETING", period="2026-1"),
        limits={"habilidad": 10},
    )
    assert curricular["habilidad"][0].document.id == "HAB_ANALIZAR"
    assert {
        item.document.id for item in curricular["habilidad"]
    }.isdisjoint({"LAB_ANALIZAR", "OTHER_CAREER"})

    laboral = retriever.retrieve(
        "Analizar indicadores",
        scope=EmbeddingScope(source_kinds=("labor",)),
        limits={"habilidad": 10},
    )
    assert [item.document.id for item in laboral["habilidad"]] == ["LAB_ANALIZAR"]


def test_limites_y_pool_configurables_y_metadatos_no_exponen_vectores() -> None:
    documentos = tuple(
        CatalogDocument(
            id=f"HAB_{indice}",
            text=f"Diseñar habilidad {indice}",
            source_kind="career_curriculum",
            career="MARKETING",
            period="2026-1",
            type="habilidad",
            catalog_version="v1",
        )
        for indice in range(5)
    )
    resultado = EmbeddingRetriever(
        _FakeEmbeddingProvider(), InMemoryEmbeddingIndex(documentos)
    ).retrieve(
        "Diseñar",
        scope=EmbeddingScope(career="MARKETING", period="2026-1"),
        limits={"habilidad": 4},
        pool_size=4,
    )

    assert len(resultado["habilidad"]) == 4
    candidato = resultado["habilidad"][0].a_dict()
    assert candidato["source_kind"] == "career_curriculum"
    assert candidato["career"] == "MARKETING"
    assert candidato["period"] == "2026-1"
    assert candidato["type"] == "habilidad"
    assert candidato["catalog_version"] == "v1"
    assert "similarity" in candidato
    assert "vector" not in candidato


def test_contexto_usa_embedding_auditable_y_fallback_lexical_si_falla() -> None:
    catalogo = CatalogoCHH(
        competencias=(),
        habilidades=(ConceptoCHH("HAB_1", "Analizar datos", "Analizar datos"),),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="lexical-v1",
    )
    perfil = {"carrera": "MARKETING", "periodo": "2026-1", "estado": "BORRADOR"}
    retriever = EmbeddingRetriever(
        _FailingEmbeddingProvider(),
        InMemoryEmbeddingIndex(
            (
                CatalogDocument(
                    id="EMB_1",
                    text="Resultado semántico",
                    source_kind="career_curriculum",
                    career="MARKETING",
                    period="2026-1",
                    type="habilidad",
                    catalog_version="embedding-v1",
                ),
            )
        ),
    )

    contexto = construir_contexto_por_logro(
        {"logro": "Analizar datos"}, catalogo, perfil, retriever=retriever
    )

    assert contexto["recuperacion"]["method"] == "lexical"
    assert contexto["recuperacion"]["minimum_similarity"] == 0.0
    assert contexto["candidatos"]["habilidad"][0]["id"] == "HAB_1"


def test_umbral_excluye_similitudes_cero_y_negativas_y_activa_fallback_lexical() -> None:
    catalogo = CatalogoCHH(
        competencias=(),
        habilidades=(ConceptoCHH("HAB_LEX", "Analizar datos", "Fallback lexical"),),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="lexical-v1",
    )
    retriever = EmbeddingRetriever(
        _VectorsProvider([[1.0, 0.0], [0.0, -1.0]]),
        InMemoryEmbeddingIndex(_documents()[:2]),
    )
    scope = EmbeddingScope.curriculum("MARKETING", "2026-1")

    with pytest.raises(EmbeddingUnavailable) as error:
        retriever.retrieve("Analizar datos", scope=scope, limits={"habilidad": 2})

    assert error.value.reason_code == FALLBACK_REASON_CANDIDATES_BELOW_THRESHOLD
    contexto = construir_contexto_por_logro(
        {"logro": "Analizar datos"},
        catalogo,
        {"carrera": "MARKETING", "periodo": "2026-1"},
        retriever=retriever,
    )
    assert contexto["recuperacion"]["method"] == "lexical"
    assert contexto["recuperacion"]["reason_code"] == FALLBACK_REASON_CANDIDATES_BELOW_THRESHOLD
    assert contexto["candidatos"]["habilidad"][0]["id"] == "HAB_LEX"


def test_auditoria_de_fallback_usa_reason_codes_estables_sin_exponer_excepciones() -> None:
    catalogo = CatalogoCHH(
        competencias=(),
        habilidades=(ConceptoCHH("HAB_LEX", "Analizar datos", "Fallback lexical"),),
        herramientas=(),
        ejemplos_por_habilidad={},
        origen=("test",),
        version="lexical-v1",
    )
    perfil = {"carrera": "MARKETING", "periodo": "2026-1"}
    base = {"logro": "Analizar datos"}

    ausente = construir_contexto_por_logro(base, catalogo, perfil)
    vacio = construir_contexto_por_logro(
        base,
        catalogo,
        perfil,
        retriever=EmbeddingRetriever(_FakeEmbeddingProvider(), InMemoryEmbeddingIndex(())),
    )
    proveedor = construir_contexto_por_logro(
        base,
        catalogo,
        perfil,
        retriever=EmbeddingRetriever(
            _FailingEmbeddingProvider(),
            InMemoryEmbeddingIndex(_documents()[:1]),
        ),
    )

    assert ausente["recuperacion"]["reason_code"] == FALLBACK_REASON_RETRIEVER_ABSENT
    assert ausente["recuperacion"]["scope"] == {
        "scope_kind": "career_curriculum",
        "career": "MARKETING",
        "period": "2026-1",
        "source_kinds": ["career_curriculum"],
    }
    assert vacio["recuperacion"]["reason_code"] == FALLBACK_REASON_CATALOG_EMPTY
    assert proveedor["recuperacion"]["reason_code"] == FALLBACK_REASON_PROVIDER_OR_VECTOR_INVALID
    serializado = json.dumps(proveedor)
    assert "provider unavailable" not in serializado
    assert "api_key" not in serializado


def test_decision_con_cita_real_pero_ajena_queda_en_revision_aunque_el_catalogo_la_sugiera(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    logro = "Analizar datos de mercado para identificar segmentos."
    id_habilidad = analista_llm._hash_id("HAB_SRC", "SIL_1", "L1", logro)
    decision = analista_llm.DecisionCurricular(
        id_habilidad_fuente=id_habilidad,
        competencia=analista_llm.ConceptoPropuesto(nombre="Gestión del talento humano"),
        habilidad=analista_llm.ConceptoPropuesto(nombre="Planificar nóminas de personal"),
        evidencia=[logro],
        confianza=0.95,
    )

    class _LLM:
        model_name = "llm-test"

        def with_structured_output(self, schema):
            self.schema = schema
            return self

        def invoke(self, _prompt: str):
            return self.schema(decisiones=[decision])

    monkeypatch.setattr(analista_llm, "obtener_llm", lambda _rol: _LLM())
    monkeypatch.setattr(
        analista_llm,
        "_cargar_perfil",
        lambda _carrera, _periodo: {"carrera": "MARKETING", "periodo": "2026-1"},
    )
    resultado = analista_llm.analizar_registros_curriculares(
        [
            {
                "id_silabo": "SIL_1",
                "datos": {
                    "curso": "Investigación de mercados",
                    "sumilla": logro,
                    "logro_general": logro,
                    "logros_especificos": [{"etiqueta": "L1", "descripcion": logro}],
                },
            }
        ],
        CatalogoCHH(
            competencias=(ConceptoCHH("COMP_TALENTO", "Gestión del talento humano", ""),),
            habilidades=(ConceptoCHH("HAB_NOMINA", "Planificar nóminas de personal", ""),),
            herramientas=(),
            ejemplos_por_habilidad={},
            origen=("test",),
            version="v1",
        ),
        "MARKETING",
        "2026-1",
        tmp_path,
        inspeccionar=False,
    )

    assert not resultado.decisiones
    reporte = next(
        fila for fila in resultado.reportes if fila.get("estado") == "REVISAR_VALIDACION"
    )
    assert "HABILIDAD_SIN_ANCLA_EVIDENCIA" in reporte["problemas"]


def test_decision_con_competencia_y_habilidad_de_raiz_compartida_requiere_cita_relevante(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    logro = "Analizar datos de mercado para identificar segmentos."
    id_habilidad = analista_llm._hash_id("HAB_SRC", "SIL_1", "L1", logro)
    decision = analista_llm.DecisionCurricular(
        id_habilidad_fuente=id_habilidad,
        competencia=analista_llm.ConceptoPropuesto(nombre="Gestión de campañas"),
        habilidad=analista_llm.ConceptoPropuesto(nombre="Gestionar campañas"),
        evidencia=[logro],
        confianza=0.95,
    )

    class _LLM:
        model_name = "llm-test"

        def with_structured_output(self, schema):
            self.schema = schema
            return self

        def invoke(self, _prompt: str):
            return self.schema(decisiones=[decision])

    monkeypatch.setattr(analista_llm, "obtener_llm", lambda _rol: _LLM())
    monkeypatch.setattr(
        analista_llm,
        "_cargar_perfil",
        lambda _carrera, _periodo: {"carrera": "MARKETING", "periodo": "2026-1"},
    )
    resultado = analista_llm.analizar_registros_curriculares(
        [
            {
                "id_silabo": "SIL_1",
                "datos": {
                    "curso": "Investigación de mercados",
                    "sumilla": logro,
                    "logro_general": logro,
                    "logros_especificos": [{"etiqueta": "L1", "descripcion": logro}],
                },
            }
        ],
        CatalogoCHH(
            competencias=(),
            habilidades=(),
            herramientas=(),
            ejemplos_por_habilidad={},
            origen=("test",),
            version="v1",
        ),
        "MARKETING",
        "2026-1",
        tmp_path,
        inspeccionar=False,
    )

    assert not resultado.decisiones
    reporte = next(
        fila for fila in resultado.reportes if fila.get("estado") == "REVISAR_VALIDACION"
    )
    assert "HABILIDAD_SIN_ANCLA_EVIDENCIA" in reporte["problemas"]


def test_auditoria_conserva_periodo_de_secuencia_variable() -> None:
    scope = analista_llm._scope_recuperacion_auditable(
        {
            "scope_kind": "career_curriculum",
            "career": "MARKETING",
            "period": "2026-123",
            "source_kinds": ["career_curriculum"],
        }
    )

    assert scope == {
        "scope_kind": "career_curriculum",
        "career": "MARKETING",
        "period": "2026-123",
        "source_kinds": ["career_curriculum"],
    }


def test_umbral_positivo_configurable_excluye_similitud_insuficiente() -> None:
    class _PositiveThresholdProvider:
        model_name = "positive-threshold-test"

        def embed_query(self, _text: str) -> list[float]:
            return [1.0, 0.0]

        def embed_documents(self, texts: list[str]) -> list[list[float]]:
            return [[0.8, 0.6] for _ in texts]

    documentos = _documents()[:1]
    scope = EmbeddingScope.curriculum("MARKETING", "2026-1")
    rechazado = EmbeddingRetriever(
        _PositiveThresholdProvider(),
        InMemoryEmbeddingIndex(documentos),
        minimum_similarity=0.81,
    )
    aceptado = EmbeddingRetriever(
        _PositiveThresholdProvider(),
        InMemoryEmbeddingIndex(documentos),
        minimum_similarity=0.75,
    )

    with pytest.raises(EmbeddingUnavailable) as error:
        rechazado.retrieve("Analizar datos", scope=scope, limits={"habilidad": 1})

    assert error.value.reason_code == FALLBACK_REASON_CANDIDATES_BELOW_THRESHOLD
    assert aceptado.retrieve("Analizar datos", scope=scope, limits={"habilidad": 1})[
        "habilidad"
    ][0].document.id == "HAB_ANALIZAR"


def test_auditoria_persistible_conserva_reason_code_por_logro_sin_secretos(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    provider = _FailingEmbeddingProvider()
    provider.config_identifier = "api_key=super-secret"  # type: ignore[attr-defined]
    retriever = EmbeddingRetriever(provider, InMemoryEmbeddingIndex(_documents()[:1]))

    class _LLM:
        model_name = "llm-test"

        def with_structured_output(self, schema):
            self.schema = schema
            return self

        def invoke(self, _prompt: str):
            return self.schema(decisiones=[])

    monkeypatch.setattr(analista_llm, "obtener_llm", lambda _rol: _LLM())
    monkeypatch.setattr(
        analista_llm,
        "_cargar_perfil",
        lambda _carrera, _periodo: {"carrera": "MARKETING", "periodo": "2026-1"},
    )
    resultado = analista_llm.analizar_registros_curriculares(
        [
            {
                "id_silabo": "SIL_1",
                "datos": {
                    "curso": "Analítica",
                    "sumilla": "Analizar datos",
                    "logro_general": "Analizar datos",
                    "logros_especificos": [
                        {"etiqueta": "L1", "descripcion": "Analizar datos"}
                    ],
                },
            }
        ],
        CatalogoCHH(
            competencias=(),
            habilidades=(ConceptoCHH("HAB_LEX", "Analizar datos", ""),),
            herramientas=(),
            ejemplos_por_habilidad={},
            origen=("test",),
            version="v1",
        ),
        "MARKETING",
        "2026-1",
        tmp_path,
        inspeccionar=False,
        embedding_retriever=retriever,
    )

    assert resultado.auditoria_contexto is not None
    recuperacion = resultado.auditoria_contexto["recuperacion_por_logro"]
    assert recuperacion[0]["reason_code"] == FALLBACK_REASON_PROVIDER_OR_VECTOR_INVALID
    assert recuperacion[0]["method"] == "lexical"
    assert recuperacion[0]["minimum_similarity"] == 0.0
    assert recuperacion[0]["scope"] == {
        "scope_kind": "career_curriculum",
        "career": "MARKETING",
        "period": "2026-1",
        "source_kinds": ["career_curriculum"],
    }
    serializado = json.dumps(resultado.reportes)
    auditoria_serializada = json.dumps(resultado.auditoria_contexto)
    assert "embedding_provider_or_vector_invalid" in serializado
    assert "super-secret" not in serializado
    assert "super-secret" not in auditoria_serializada
    assert "provider unavailable" not in serializado


def test_analista_consulta_un_embedding_por_logro_y_envia_su_contexto(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    provider = _FakeEmbeddingProvider()
    retriever = EmbeddingRetriever(
        provider,
        InMemoryEmbeddingIndex(
            (
                CatalogDocument(
                    id="EMB_ANALIZAR",
                    text="Contexto recuperado para analizar",
                    source_kind="career_curriculum",
                    career="MARKETING",
                    period="2026-1",
                    type="habilidad",
                    catalog_version="v1",
                ),
                CatalogDocument(
                    id="EMB_DISEÑAR",
                    text="Contexto recuperado para diseñar",
                    source_kind="career_curriculum",
                    career="MARKETING",
                    period="2026-1",
                    type="habilidad",
                    catalog_version="v1",
                ),
            )
        ),
    )
    prompts: list[str] = []

    class _LLM:
        model_name = "llm-test"

        def with_structured_output(self, schema):
            self.schema = schema
            return self

        def invoke(self, prompt: str):
            prompts.append(prompt)
            return self.schema(decisiones=[])

    monkeypatch.setattr(analista_llm, "obtener_llm", lambda _rol: _LLM())
    monkeypatch.setattr(
        analista_llm,
        "_cargar_perfil",
        lambda _carrera, _periodo: {
            "carrera": "MARKETING",
            "periodo": "2026-1",
            "estado": "BORRADOR",
        },
    )
    registros = [
        {
            "id_silabo": "SIL_1",
            "datos": {
                "curso": "Curso",
                "sumilla": "Sumilla",
                "logro_general": "Logro general",
                "logros_especificos": [
                    {"etiqueta": "L1", "descripcion": "Analizar datos"},
                    {"etiqueta": "L2", "descripcion": "Diseñar campañas"},
                ],
            },
        }
    ]

    analista_llm.analizar_registros_curriculares(
        registros,
        CatalogoCHH((), (), (), {}, ("test",), "v1"),
        "MARKETING",
        "2026-1",
        tmp_path,
        inspeccionar=False,
        embedding_retriever=retriever,
    )

    assert len(provider.queries) == 2
    assert provider.queries[0] != provider.queries[1]
    assert "Analizar datos" in provider.queries[0]
    assert "Diseñar campañas" in provider.queries[1]
    assert all("Curso" in query for query in provider.queries)
    assert prompts
    serializados = "\n".join(prompts)
    assert "Contexto recuperado para analizar" in serializados
    assert "Contexto recuperado para diseñar" in serializados


def test_retriever_exige_scope_explicito_y_modela_labor_como_global() -> None:
    documentos = (
        CatalogDocument(
            id="LAB_1",
            text="Analizar requisitos laborales",
            source_kind="labor",
            career=None,
            period=None,
            type="habilidad",
            catalog_version="labor-v1",
        ),
    )
    retriever = EmbeddingRetriever(_FakeEmbeddingProvider(), InMemoryEmbeddingIndex(documentos))

    with pytest.raises(EmbeddingUnavailable, match="explicit scope"):
        retriever.retrieve("Analizar")

    laboral = retriever.retrieve(
        "Analizar",
        scope=EmbeddingScope.labor_global(),
        limits={"habilidad": 1},
    )
    assert [item.document.id for item in laboral["habilidad"]] == ["LAB_1"]


@pytest.mark.parametrize(
    ("vectors", "error"),
    [
        ([[0.0, 0.0]], "zero norm"),
        ([[1.0, 0.0, 0.0]], "dimensions"),
    ],
)
def test_retriever_convierte_vectores_invalidos_en_fallback_recuperable(
    vectors: list[list[float]], error: str
) -> None:
    retriever = EmbeddingRetriever(
        _VectorsProvider(vectors),
        InMemoryEmbeddingIndex(_documents()[:1]),
    )

    with pytest.raises(EmbeddingUnavailable, match=error):
        retriever.retrieve(
            "Analizar",
            scope=EmbeddingScope.curriculum("MARKETING", "2026-1"),
            limits={"habilidad": 1},
        )


def test_indice_no_guarda_vectores_parciales_si_un_vector_de_documento_es_invalido() -> None:
    class _Provider(_FakeEmbeddingProvider):
        def __init__(self) -> None:
            super().__init__()
            self.calls = 0

        def embed_documents(self, texts: list[str]) -> list[list[float]]:
            self.calls += 1
            return [[1.0, 0.0], [0.0, 0.0]]

    provider = _Provider()
    indice = InMemoryEmbeddingIndex(_documents()[:2])
    retriever = EmbeddingRetriever(provider, indice)
    scope = EmbeddingScope.curriculum("MARKETING", "2026-1")

    with pytest.raises(EmbeddingUnavailable, match="zero norm"):
        retriever.retrieve("Analizar", scope=scope, limits={"habilidad": 2})
    with pytest.raises(EmbeddingUnavailable, match="zero norm"):
        retriever.retrieve("Analizar", scope=scope, limits={"habilidad": 2})

    assert provider.calls == 2


def test_documento_rechaza_scope_de_fuente_desconocido() -> None:
    with pytest.raises(ValueError, match="unsupported embedding source scope"):
        CatalogDocument(
            id="UNKNOWN_1",
            text="Texto",
            source_kind="unknown",  # type: ignore[arg-type]
            career=None,
            period=None,
            type="habilidad",
            catalog_version="v1",
        )


def test_retriever_falla_cuando_el_catalogo_embedding_elegible_esta_vacio() -> None:
    retriever = EmbeddingRetriever(_FakeEmbeddingProvider(), InMemoryEmbeddingIndex(()))

    with pytest.raises(EmbeddingUnavailable, match="empty"):
        retriever.retrieve(
            "Analizar",
            scope=EmbeddingScope.curriculum("MARKETING", "2026-1"),
            limits={"habilidad": 1},
        )


def test_indice_no_colisiona_ids_repetidos_de_scopes_o_versiones_y_rechaza_identidades_iguales(
) -> None:
    marketing = CatalogDocument(
        id="HAB_1",
        text="Analizar marketing",
        source_kind="career_curriculum",
        career="MARKETING",
        period="2026-1",
        type="habilidad",
        catalog_version="v1",
    )
    ingenieria = CatalogDocument(
        id="HAB_1",
        text="Analizar ingeniería",
        source_kind="career_curriculum",
        career="INGENIERIA",
        period="2026-1",
        type="habilidad",
        catalog_version="v2",
    )
    indice = InMemoryEmbeddingIndex((marketing, ingenieria))
    retriever = EmbeddingRetriever(_FakeEmbeddingProvider(), indice)

    assert retriever.retrieve(
        "Analizar", scope=EmbeddingScope.curriculum("MARKETING", "2026-1"), limits={"habilidad": 1}
    )["habilidad"][0].document.text == "Analizar marketing"
    assert retriever.retrieve(
        "Analizar", scope=EmbeddingScope.curriculum("INGENIERIA", "2026-1"), limits={"habilidad": 1}
    )["habilidad"][0].document.text == "Analizar ingeniería"

    with pytest.raises(ValueError, match="duplicate embedding document identity"):
        InMemoryEmbeddingIndex((marketing, marketing))


def test_limites_invalidos_no_se_silencian_y_mapping_vacio_es_sin_candidatos() -> None:
    catalogo = CatalogoCHH(
        (),
        (ConceptoCHH("HAB_1", "Analizar datos", "Analizar datos"),),
        (),
        {},
        ("test",),
        "v1",
    )
    perfil = {"carrera": "MARKETING", "periodo": "2026-1"}

    vacio = construir_contexto_por_logro(
        {"logro": "Analizar"}, catalogo, perfil, limites_candidatos={}
    )
    assert vacio["candidatos"] == {"competencia": [], "habilidad": [], "herramienta": []}

    with pytest.raises(ValueError, match="non-negative"):
        construir_contexto_por_logro(
            {"logro": "Analizar"}, catalogo, perfil, limites_candidatos={"habilidad": -1}
        )


def test_pool_es_por_tipo_antes_del_limite_y_el_audit_no_filtra_secretos() -> None:
    provider = _FakeEmbeddingProvider()
    provider.config_identifier = "api_key=super-secret"  # type: ignore[attr-defined]
    retriever = EmbeddingRetriever(provider, InMemoryEmbeddingIndex(_documents()))
    catalogo = CatalogoCHH((), (), (), {}, ("test",), "v1")
    contexto = construir_contexto_por_logro(
        {"logro": "Analizar datos"},
        catalogo,
        {"carrera": "MARKETING", "periodo": "2026-1"},
        retriever=retriever,
        limites_candidatos={"habilidad": 4},
        pool_retrieval=1,
    )

    assert len(contexto["candidatos"]["habilidad"]) == 1
    serializado = json.dumps(contexto)
    assert "super-secret" not in serializado
    assert "vector" not in serializado
    assert contexto["recuperacion"]["config"].startswith("provider:")

    with pytest.raises(ValueError, match="pool_size"):
        retriever.retrieve(
            "Analizar",
            scope=EmbeddingScope.curriculum("MARKETING", "2026-1"),
            pool_size=-1,
        )


def test_nueve_logros_mantienen_retrieval_independiente_y_lotes_llm_de_ocho(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    provider = _FakeEmbeddingProvider()
    retriever = EmbeddingRetriever(provider, InMemoryEmbeddingIndex(_documents()[:2]))
    prompts: list[str] = []

    class _LLM:
        model_name = "llm-test"

        def with_structured_output(self, schema):
            self.schema = schema
            return self

        def invoke(self, prompt: str):
            prompts.append(prompt)
            return self.schema(decisiones=[])

    monkeypatch.setattr(analista_llm, "obtener_llm", lambda _rol: _LLM())
    monkeypatch.setattr(
        analista_llm,
        "_cargar_perfil",
        lambda _carrera, _periodo: {"carrera": "MARKETING", "periodo": "2026-1"},
    )
    logros = [
        {"etiqueta": f"L{indice}", "descripcion": f"Analizar datos {indice}"}
        for indice in range(9)
    ]
    analista_llm.analizar_registros_curriculares(
        [{"id_silabo": "SIL_1", "datos": {"curso": "Curso", "logros_especificos": logros}}],
        CatalogoCHH((), (), (), {}, ("test",), "v1"),
        "MARKETING",
        "2026-1",
        tmp_path,
        inspeccionar=False,
        embedding_retriever=retriever,
    )

    assert len(provider.queries) == 9
    assert len(set(provider.queries)) == 9
    tamanos_lote = sorted(prompt.count('"id_habilidad_fuente"') for prompt in prompts)
    assert tamanos_lote == [1, 1, 8, 8]


def test_limpieza_no_activa_embeddings_para_carrera_sin_allowlist(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    fuente = tmp_path / "silabo.docx"
    documento = Document()
    documento.add_paragraph("Sílabo de prueba")
    documento.save(fuente)
    validacion = validar_archivo(fuente, "Marketing", "2026-1")
    catalogo = CatalogoCHH(
        (),
        (ConceptoCHH("HAB_1", "Analizar datos", "Analizar datos"),),
        (),
        {},
        ("test",),
        "v1",
    )
    capturados: list[object] = []
    monkeypatch.setenv("NORMALIZADOR_CURRICULAR_EMBEDDINGS", "true")
    monkeypatch.setattr(
        limpieza_silabos,
        "analizar_registros_curriculares",
        lambda *_args, **kwargs: (
            capturados.append(kwargs["embedding_retriever"])
            or ResultadoAnalisisCurricular({}, (), "fake", "no_ejecutado", 0)
        ),
    )

    limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        catalogo,
        usar_llm=True,
        embedding_provider=_FakeEmbeddingProvider(),
    )

    assert capturados[0] is None


def test_limpieza_indexa_solo_el_catalogo_curricular_de_carrera(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    fuente = tmp_path / "silabo.docx"
    documento = Document()
    documento.add_paragraph("Sílabo de prueba")
    documento.save(fuente)
    validacion = validar_archivo(fuente, "INGENIERIA", "2026-1")
    catalogo_global = CatalogoCHH(
        (),
        (ConceptoCHH("LAB_ONLY", "Requerimiento laboral", "labor"),),
        (),
        {},
        ("global",),
        "global-v1",
    )
    catalogo_carrera = CatalogoCHH(
        (),
        (ConceptoCHH("CURR_ONLY", "Habilidad curricular", "curricular"),),
        (),
        {},
        ("career",),
        "career-v1",
    )
    capturados: list[object] = []
    monkeypatch.setenv("NORMALIZADOR_CURRICULAR_EMBEDDINGS", "true")
    monkeypatch.setenv("NORMALIZADOR_CURRICULAR_EMBEDDING_CARRERAS", "INGENIERIA@2026-1")
    monkeypatch.setattr(limpieza_silabos, "cargar_catalogo", lambda: catalogo_global)
    monkeypatch.setattr(
        limpieza_silabos,
        "cargar_catalogo_carrera",
        lambda _carrera, _periodo: catalogo_carrera,
    )
    monkeypatch.setattr(
        limpieza_silabos,
        "analizar_registros_curriculares",
        lambda *_args, **kwargs: (
            capturados.append(kwargs["embedding_retriever"])
            or ResultadoAnalisisCurricular({}, (), "fake", "no_ejecutado", 0)
        ),
    )

    limpiar_archivo(
        fuente,
        tmp_path / "ejecucion",
        validacion,
        usar_llm=True,
        embedding_provider=_FakeEmbeddingProvider(),
    )

    retriever = capturados[0]
    assert isinstance(retriever, EmbeddingRetriever)
    assert [document.id for document in retriever.index.documents] == ["CURR_ONLY"]


@pytest.mark.parametrize(
    ("allowlist", "periodo", "esperado"),
    [
        ("Ingeniería de Sistemas@2026-1", "2026-1", True),
        ("Ingeniería de Sistemas@2025-1", "2026-1", False),
        ("Ingeniería de Sistemas", "2026-1", False),
    ],
)
def test_allowlist_curricular_exige_pareja_carrera_periodo(
    monkeypatch: pytest.MonkeyPatch,
    allowlist: str,
    periodo: str,
    esperado: bool,
) -> None:
    monkeypatch.setenv("NORMALIZADOR_CURRICULAR_EMBEDDING_CARRERAS", allowlist)

    assert (
        limpieza_silabos._embeddings_curriculares_habilitados(
            "Ingeniería de Sistemas", periodo, True
        )
        is esperado
    )


def test_factory_opt_in_sin_credenciales_deja_el_fallback_lexical(
    monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    catalogo = CatalogoCHH((), (), (), {}, ("test",), "v1")

    assert (
        crear_retriever_curricular_opt_in(
            catalogo,
            career="MARKETING",
            period="2026-1",
            enabled=True,
        )
        is None
    )
