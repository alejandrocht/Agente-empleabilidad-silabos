"""Limpieza y extracción estructural de DOCX/PDF curriculares."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from collections.abc import Callable
from dataclasses import replace
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from agente.normalizador.empleabilidad.catalogo import (
    CatalogoCHH,
    cargar_catalogo,
    cargar_catalogo_carrera,
)
from agente.normalizador.empleabilidad.entrada import normalizar_etiqueta
from agente.normalizador.modelos import (
    ArchivoSilabo,
    Hallazgo,
    ProgresoLimpiezaLLM,
    ResultadoLimpiezaSilabos,
    ResultadoValidacionSilabos,
)
from agente.normalizador.silabos.analista_llm import analizar_registros_curriculares
from agente.normalizador.silabos.salida import (
    ResultadoCatalogoCurricular,
    _catalogo_curricular,
    construir_salidas_curriculares,
)

_PATRON_CODIGO_CURRICULAR = r"(?<![A-Z0-9])(?:[GE]\d+|[GE]{2,4})(?![A-Z0-9])"
_SECCIONES_HERRAMIENTAS = (
    "recursos de aprendizaje",
    "recursos tecnologicos",
    "recursos tecnológicos",
    "software",
    "herramientas digitales",
)


def _texto(valor: object) -> str:
    return re.sub(r"\s+", " ", str(valor or "")).strip()


def _clave(valor: object) -> str:
    return normalizar_etiqueta(valor).replace(" ", "_")


def _hash_id(prefijo: str, *partes: str) -> str:
    payload = "|".join(partes).encode("utf-8")
    return f"{prefijo}_{hashlib.sha256(payload).hexdigest()[:16]}"


def _hallazgo(
    codigo: str,
    severidad: str,
    mensaje: str,
    archivo: str,
    detalle: str | None = None,
) -> Hallazgo:
    return Hallazgo(
        codigo=codigo,
        severidad=severidad,  # type: ignore[arg-type]
        mensaje=mensaje,
        hoja=archivo,
        detalle=detalle,
    )


def _filas_tabla(tabla: Any) -> list[list[str]]:
    """Devuelve filas lógicas sin repetir continuaciones de ``w:vMerge``.

    ``python-docx`` proyecta el texto de la celda que inicia una combinación
    vertical sobre cada fila de continuación. Cuando toda la fila continúa una
    combinación, no representa un registro curricular adicional.
    """

    filas: list[list[str]] = []
    for fila in tabla.rows:
        celdas_fisicas = fila._tr.tc_lst
        if celdas_fisicas and all(_es_continuacion_vertical(celda) for celda in celdas_fisicas):
            continue
        filas.append([_texto(celda.text) for celda in fila.cells])
    return filas


def _es_continuacion_vertical(celda: Any) -> bool:
    combinacion = celda.tcPr.vMerge
    return combinacion is not None and combinacion.val != "restart"


def _codigos(texto: str, aceptar_no_numericos: bool = True) -> list[str]:
    """Extrae códigos declarados, incluidos formatos curriculares no estándar."""

    valor = _texto(texto).upper()
    numericos = re.findall(r"(?<![A-Z0-9])[GE]\d+(?![A-Z0-9])", valor)
    if not aceptar_no_numericos:
        return list(dict.fromkeys(numericos))
    if numericos:
        palabras_cortas = re.findall(
            r"(?<![A-Z0-9])[A-Z]{2,4}(?![A-Z0-9])",
            valor,
        )
        if len(valor.split()) <= 4:
            return list(dict.fromkeys(numericos + palabras_cortas))
        return list(dict.fromkeys(numericos))

    # Algunas carreras usan referencias como ``EE``. Solo se aceptan códigos
    # alfabéticos cuando la celda es corta; así no se convierten palabras de una
    # descripción extensa en competencias.
    palabras = re.findall(r"(?<![A-Z0-9])[A-Z]{2,4}(?![A-Z0-9])", valor)
    if len(valor.split()) <= 4:
        return list(dict.fromkeys(palabras))
    return []


def _extraer_docx(
    ruta: Path,
    nombre: str,
    carrera: str,
    periodo: str,
) -> dict[str, object]:
    doc = Document(str(ruta))
    metadata: dict[str, str] = {}
    if doc.tables:
        for fila in doc.tables[0].rows:
            valores = [_texto(celda.text) for celda in fila.cells]
            if len(valores) >= 2 and valores[0]:
                metadata[_clave(valores[0])] = valores[1]

    sumilla = ""
    if len(doc.tables) > 1:
        candidato = _texto(
            " ".join(celda.text for fila in doc.tables[1].rows for celda in fila.cells)
        )
        if candidato and _clave(candidato) != "sumilla":
            sumilla = candidato
    logro_general = ""
    logros: list[dict[str, object]] = []
    competencias: list[dict[str, str]] = []
    programa: list[str] = []
    herramientas_evidencia: list[dict[str, str]] = []

    for tabla in doc.tables:
        filas = _filas_tabla(tabla)
        if not filas:
            continue
        encabezado = _clave(" ".join(filas[0]))
        if "competencias_genericas" in encabezado or "competencias_especificas" in encabezado:
            for valores in filas[1:]:
                if len(valores) < 3:
                    continue
                encontrados = _codigos(valores[-1])
                if encontrados and valores[0] and not re.fullmatch(r"L\d+", valores[0], re.I):
                    for codigo in encontrados:
                        competencias.append(
                            {
                                "codigo": codigo,
                                "nombre": valores[0].strip(" ."),
                                "descripcion": valores[1].strip(),
                            }
                        )
        if "logro_de_aprendizaje_general" in encabezado:
            if len(filas) > 1 and filas[1]:
                logro_general = filas[1][0]
            for valores in filas[1:]:
                if len(valores) >= 3 and re.fullmatch(r"L\d+", valores[0], re.I):
                    logros.append(
                        {
                            "etiqueta": valores[0].upper(),
                            "descripcion": valores[1],
                            "codigos_competencia": _codigos(valores[-1]),
                        }
                    )
        if "semana" in encabezado and ("tema" in encabezado or "contenido" in encabezado):
            for valores in filas[1:]:
                if valores and re.fullmatch(r"(?:[1-9]|1[0-5])", valores[0]):
                    programa.append(" | ".join(valor for valor in valores[1:3] if valor))
        herramientas_evidencia.extend(_herramientas_desde_tabla(filas))

    herramientas_evidencia.extend(_herramientas_desde_parrafos(doc))

    if not sumilla:
        for tabla in doc.tables:
            filas = _filas_tabla(tabla)
            for indice, valores in enumerate(filas):
                if valores and _clave(valores[0]) == "sumilla" and indice + 1 < len(filas):
                    sumilla = filas[indice + 1][0]
                    break
            if sumilla:
                break

    curso = _primer_metadata(metadata, ("nombre_del_curso", "curso", "asignatura"))
    curso = curso or _nombre_desde_archivo(nombre)
    ciclo = _ciclo_desde_ruta(nombre) or _primer_metadata(metadata, ("ciclo", "nivel"))
    texto_relevante = " ".join(
        parte
        for parte in (
            sumilla,
            logro_general,
            " ".join(str(logro["descripcion"]) for logro in logros),
            " ".join(programa),
        )
        if parte
    )
    texto_fuente = _texto(
        " ".join(
            [parrafo.text for parrafo in doc.paragraphs]
            + [celda.text for tabla in doc.tables for fila in tabla.rows for celda in fila.cells]
        )
    )
    id_silabo = _hash_id("SIL", carrera, periodo, nombre)
    id_curso = _hash_id("CUR", carrera, periodo, curso)
    return {
        "id_silabo": id_silabo,
        "id_curso": id_curso,
        "carrera": carrera,
        "periodo": periodo,
        "origen": {"archivo": nombre, "formato": "docx"},
        "datos": {
            "curso": curso,
            "ciclo": ciclo,
            "codigo_curso": _primer_metadata(metadata, ("codigo_del_curso", "codigo")),
            "sumilla": sumilla,
            "logro_general": logro_general,
            "logros_especificos": logros,
            "competencias_declaradas": competencias,
            "programa_analitico": programa,
            "herramientas_evidencia": _deduplicar_evidencias_herramientas(herramientas_evidencia),
            "texto_relevante": texto_relevante,
            "texto_fuente": texto_fuente,
        },
    }


def _extraer_pdf(
    ruta: Path,
    nombre: str,
    carrera: str,
    periodo: str,
) -> dict[str, object]:
    lector = PdfReader(ruta)
    texto_fuente = "\n".join(page.extract_text() or "" for page in lector.pages)
    texto = _texto(texto_fuente)
    curso = _campo_pdf(texto_fuente, r"Asignatura") or _nombre_desde_archivo(nombre)
    codigo_curso = _campo_pdf(texto_fuente, r"Código")
    sumilla = _seccion_pdf(texto_fuente, r"II\.\s*Sumilla", r"III\.\s*Competencias")
    bloque_competencias = _seccion_pdf(
        texto_fuente,
        r"III\.\s*Competencias",
        r"IV\.\s*Logros de aprendizaje",
    )
    bloque_logros = _seccion_pdf(
        texto_fuente,
        r"IV\.\s*Logros de aprendizaje",
        r"V\.\s*Estrategia de enseñanza",
    )
    competencias = _competencias_pdf(bloque_competencias)
    logro_general = _logro_general_pdf(bloque_logros)
    logros = _logros_pdf(bloque_logros)
    id_silabo = _hash_id("SIL", carrera, periodo, nombre)
    id_curso = _hash_id("CUR", carrera, periodo, curso)
    return {
        "id_silabo": id_silabo,
        "id_curso": id_curso,
        "carrera": carrera,
        "periodo": periodo,
        "origen": {"archivo": nombre, "formato": "pdf"},
        "datos": {
            "curso": curso,
            "ciclo": _ciclo_pdf(texto_fuente) or _ciclo_desde_ruta(nombre),
            "codigo_curso": codigo_curso,
            "sumilla": sumilla,
            "logro_general": logro_general,
            "logros_especificos": logros,
            "competencias_declaradas": competencias,
            "programa_analitico": [],
            "herramientas_evidencia": _evidencias_herramientas_pdf(texto_fuente),
            "texto_relevante": texto,
            "texto_fuente": texto,
        },
    }


def _seccion_pdf(texto: str, inicio: str, fin: str) -> str:
    patron = rf"(?is){inicio}\s*(.*?)(?={fin})"
    coincidencia = re.search(patron, texto)
    return _texto(coincidencia.group(1)) if coincidencia else ""


def _campo_pdf(texto: str, etiqueta: str) -> str:
    coincidencia = re.search(rf"(?im)^\s*{etiqueta}\s+([^\n]+)", texto)
    return _texto(coincidencia.group(1)) if coincidencia else ""


def _competencias_pdf(texto: str) -> list[dict[str, str]]:
    cuerpo = re.sub(r"(?i)^competencias\s+(?:genéricas|específicas)\s+", "", texto)
    coincidencias = list(re.finditer(_PATRON_CODIGO_CURRICULAR, cuerpo, flags=re.IGNORECASE))
    resultado: list[dict[str, str]] = []
    for indice, coincidencia in enumerate(coincidencias):
        inicio = coincidencias[indice - 1].end() if indice else 0
        segmento = _texto(cuerpo[inicio : coincidencia.start()])
        segmento = re.sub(r"(?i)^competencias\s+(?:genéricas|específicas)\s+", "", segmento)
        if not segmento:
            continue
        verbos = (
            "Analiza|Analizar|Aplica|Aplicar|Argumenta|Desarrolla|Desarrollar|"
            "Diseña|Diseñar|Evalúa|Evaluar|Gestiona|Gestionar|Interpreta|"
            "Interpretar|Planifica|Planificar|Propone|Proponer|Reconoce|"
            "Reconocer|Utiliza|Utilizar"
        )
        verbo = re.search(rf"\b(?:{verbos})\b", segmento, flags=re.IGNORECASE)
        if verbo:
            nombre = _texto(segmento[: verbo.start()]).strip(" -")
            descripcion = _texto(segmento[verbo.start() :])
        else:
            palabras = segmento.split()
            nombre = " ".join(palabras[: min(6, len(palabras))])
            descripcion = segmento
        if nombre and descripcion:
            resultado.append(
                {
                    "codigo": coincidencia.group(0).upper(),
                    "nombre": nombre,
                    "descripcion": descripcion,
                }
            )
    return resultado


def _logro_general_pdf(texto: str) -> str:
    coincidencia = re.search(
        r"(?is)Logro de aprendizaje general\s+(.*?)(?=Logros de aprendizaje específicos)",
        texto,
    )
    return _texto(coincidencia.group(1)) if coincidencia else ""


def _logros_pdf(texto: str) -> list[dict[str, object]]:
    inicio = re.search(r"(?is)Logros de aprendizaje específicos\s+", texto)
    if inicio is None:
        return []
    cuerpo = texto[inicio.end() :]
    marcas = list(re.finditer(r"\bL\d+\b", cuerpo, flags=re.IGNORECASE))
    resultado: list[dict[str, object]] = []
    for indice, marca in enumerate(marcas):
        fin = marcas[indice + 1].start() if indice + 1 < len(marcas) else len(cuerpo)
        segmento = cuerpo[marca.end() : fin]
        codigos = _codigos_curriculares_pdf(segmento)
        descripcion = _texto(re.sub(_PATRON_CODIGO_CURRICULAR, " ", segmento, flags=re.IGNORECASE))
        if descripcion:
            resultado.append(
                {
                    "etiqueta": marca.group(0).upper(),
                    "descripcion": descripcion,
                    "codigos_competencia": codigos,
                }
            )
    return resultado


def _codigos_curriculares_pdf(texto: str) -> list[str]:
    """Reconoce E/G numéricos y alfabéticos sin convertir palabras en códigos."""

    return list(
        dict.fromkeys(
            coincidencia.group(0).upper()
            for coincidencia in re.finditer(
                _PATRON_CODIGO_CURRICULAR,
                _texto(texto).upper(),
            )
        )
    )


def _seccion_herramientas(encabezado: str) -> str:
    normalizado = encabezado.replace("_", " ")
    return next(
        (seccion for seccion in _SECCIONES_HERRAMIENTAS if seccion == normalizado),
        "",
    )


def _herramientas_desde_tabla(filas: list[list[str]]) -> list[dict[str, str]]:
    """Toma solo la fila que sigue a un encabezado estructurado exacto."""

    resultado: list[dict[str, str]] = []
    for indice, fila in enumerate(filas[:-1]):
        seccion = _seccion_herramientas(_clave(" ".join(fila)))
        if not seccion:
            continue
        evidencia = _texto(" ".join(filas[indice + 1]))
        if evidencia:
            resultado.append({"seccion": seccion, "texto": evidencia})
    return resultado


def _herramientas_desde_parrafos(doc: Any) -> list[dict[str, str]]:
    """Extrae solo el contenido inmediatamente posterior a un encabezado confiable."""

    resultado: list[dict[str, str]] = []
    parrafos = [_texto(parrafo.text) for parrafo in doc.paragraphs]
    for indice, parrafo in enumerate(parrafos):
        seccion = _seccion_herramientas(_clave(parrafo))
        if not seccion:
            continue
        for candidato in parrafos[indice + 1 : indice + 3]:
            if not candidato or _seccion_herramientas(_clave(candidato)):
                break
            if re.match(r"^(?:[IVXLC]+|\d+)\s*[.)]", candidato, flags=re.IGNORECASE):
                break
            resultado.append({"seccion": seccion, "texto": candidato})
            break
    return resultado


def _evidencias_herramientas_pdf(texto: str) -> list[dict[str, str]]:
    """Obtiene secciones PDF tituladas explícitamente como recursos o software."""

    patron_titulo = "|".join(re.escape(seccion) for seccion in _SECCIONES_HERRAMIENTAS)
    patron = re.compile(
        rf"(?ims)^\s*(?:[IVXLC]+|\d+)\s*[.)-]?\s*({patron_titulo})\s*[:.\-]?\s*"
        rf"(.*?)(?=^\s*(?:[IVXLC]+|\d+)\s*[.)-]\s+|\Z)"
    )
    resultado: list[dict[str, str]] = []
    for coincidencia in patron.finditer(texto):
        evidencia = _texto(coincidencia.group(2))
        if evidencia:
            resultado.append({"seccion": _texto(coincidencia.group(1)).lower(), "texto": evidencia})
    return _deduplicar_evidencias_herramientas(resultado)


def _deduplicar_evidencias_herramientas(
    evidencias: list[dict[str, str]],
) -> list[dict[str, str]]:
    resultado: list[dict[str, str]] = []
    vistos: set[tuple[str, str]] = set()
    for evidencia in evidencias:
        seccion = _texto(evidencia.get("seccion"))
        texto = _texto(evidencia.get("texto"))
        clave = (seccion.lower(), texto.lower())
        if seccion and texto and clave not in vistos:
            vistos.add(clave)
            resultado.append({"seccion": seccion, "texto": texto})
    return resultado


def _ciclo_pdf(texto: str) -> str:
    niveles = {
        "primero": "01",
        "segundo": "02",
        "tercero": "03",
        "cuarto": "04",
        "quinto": "05",
        "sexto": "06",
        "séptimo": "07",
        "septimo": "07",
        "octavo": "08",
        "noveno": "09",
        "décimo": "10",
        "decimo": "10",
    }
    nivel = _campo_pdf(texto, r"Nivel").lower()
    return niveles.get(nivel, nivel)


def limpiar_archivo(
    ruta_entrada: Path,
    directorio_ejecucion: Path,
    validacion: ResultadoValidacionSilabos,
    catalogo: CatalogoCHH | None = None,
    usar_llm: bool = False,
    inspeccionar_llm: bool = True,
    al_actualizar_progreso_llm: Callable[[ProgresoLimpiezaLLM], None] | None = None,
    progreso_inicial: ProgresoLimpiezaLLM | None = None,
    id_ejecucion: str = "",
) -> ResultadoLimpiezaSilabos:
    """Materializa la fuente y construye el paquete curricular CSV.

    La entrada de producción del normalizador curricular activa ``usar_llm``;
    el parámetro explícito se conserva para pruebas offline y seams controlados.
    Cuando está activo, el LLM decide la normalización semántica por lotes y
    Python conserva la evidencia, IDs, esquema y relaciones.
    """

    fuentes = directorio_ejecucion / "fuentes_curriculares"
    limpios = directorio_ejecucion / "limpios"
    reportes = directorio_ejecucion / "salidas" / "reportes"
    fuentes.mkdir(parents=True, exist_ok=True)
    limpios.mkdir(parents=True, exist_ok=True)
    reportes.mkdir(parents=True, exist_ok=True)
    cuarentena: list[dict[str, object]] = []
    hallazgos: list[Hallazgo] = []
    registros: list[dict[str, object]] = []
    progreso_extraccion = progreso_inicial or ProgresoLimpiezaLLM(
        fase="preparando",
        chunks_completados=0,
        chunks_totales=0,
        logros_procesados=0,
        logros_totales=0,
        silabos_procesados=0,
        silabos_totales=len(validacion.archivos),
        decisiones_cacheadas=0,
        reintentos=0,
        silabos_detectados=0,
        mensaje="Preparando la extracción de sílabos.",
    ).con_evento("Preparando la extracción de sílabos.")

    def publicar_progreso(progreso: ProgresoLimpiezaLLM) -> None:
        nonlocal progreso_extraccion
        progreso_extraccion = progreso
        if al_actualizar_progreso_llm is not None:
            al_actualizar_progreso_llm(progreso)

    if usar_llm and progreso_inicial is None:
        publicar_progreso(progreso_extraccion)

    materializados = _materializar(ruta_entrada, fuentes, validacion.archivos)
    for indice_archivo, archivo in enumerate(validacion.archivos, start=1):
        ruta = materializados[archivo.nombre]
        logros_archivo = 0
        silabo_extraido = False
        try:
            if archivo.formato == "docx":
                registro = _extraer_docx(
                    ruta,
                    archivo.nombre,
                    validacion.carrera,
                    validacion.periodo,
                )
            else:
                registro = _extraer_pdf(
                    ruta,
                    archivo.nombre,
                    validacion.carrera,
                    validacion.periodo,
                )
            datos = registro["datos"]
            if isinstance(datos, dict) and not datos.get("texto_relevante"):
                hallazgo = _hallazgo(
                    "SILABO_SIN_TEXTO_RELEVANTE",
                    "warning",
                    "El archivo no produjo texto curricular utilizable.",
                    archivo.nombre,
                )
                hallazgos.append(hallazgo)
                cuarentena.append(
                    {
                        "id_silabo": registro["id_silabo"],
                        "origen": registro["origen"],
                        "codigo": hallazgo.codigo,
                        "mensaje": hallazgo.mensaje,
                    }
                )
            registros.append(registro)
            silabo_extraido = True
            if isinstance(datos, dict):
                logros_archivo = len(datos.get("logros_especificos", []))
        except Exception as exc:
            hallazgo = _hallazgo(
                "SILABO_ILEGIBLE",
                "error",
                "No se pudo extraer la estructura del sílabo.",
                archivo.nombre,
                f"{type(exc).__name__}: {str(exc)[:200]}",
            )
            hallazgos.append(hallazgo)
            cuarentena.append(
                {
                    "id_archivo": archivo.nombre,
                    "codigo": hallazgo.codigo,
                    "mensaje": hallazgo.mensaje,
                    "detalle": hallazgo.detalle,
                }
            )
        if usar_llm:
            logros_detectados = 0
            for registro_extraido in registros:
                datos_extraidos = registro_extraido.get("datos")
                if not isinstance(datos_extraidos, dict):
                    continue
                logros_extraidos = datos_extraidos.get("logros_especificos", [])
                if isinstance(logros_extraidos, list):
                    logros_detectados += len(logros_extraidos)
            silabos_detectados = len(
                {
                    str(registro.get("id_silabo") or "")
                    for registro in registros
                    if registro.get("id_silabo")
                }
            )
            progreso_extraccion = replace(
                progreso_extraccion,
                fase="extrayendo",
                logros_detectados=logros_detectados,
                logros_totales=logros_detectados,
                silabos_detectados=silabos_detectados,
                silabos_procesados=0,
                silabos_totales=len(validacion.archivos),
            ).con_evento(
                f"Logros detectados: {logros_detectados}. Sílabos detectados: "
                f"{silabos_detectados}/{len(validacion.archivos)}.",
                logros_chunk=logros_archivo,
                silabos_chunk=1 if silabo_extraido else 0,
            )
            publicar_progreso(progreso_extraccion)

    staging = limpios / "silabos.jsonl"
    with staging.open("w", encoding="utf-8", newline="\n") as salida_staging:
        for registro in registros:
            salida_staging.write(
                json.dumps(registro, ensure_ascii=False, separators=(",", ":")) + "\n"
            )
    try:
        catalogo_base = catalogo or cargar_catalogo()
        # Un catálogo inyectado representa el contexto completo de la
        # ejecución (pruebas o ejecución controlada), por lo que no debe
        # mezclarse silenciosamente con un perfil instalado en disco.
        catalogo_carrera = (
            None
            if catalogo is not None
            else cargar_catalogo_carrera(
                validacion.carrera,
                validacion.periodo,
            )
        )
        decisiones_llm = {}
        analisis_llm = None
        if usar_llm:
            try:
                catalogo_para_llm = _catalogo_curricular(
                    registros,
                    catalogo_base,
                    catalogo_carrera,
                )
                analisis_llm = analizar_registros_curriculares(
                    registros,
                    catalogo_para_llm,
                    validacion.carrera,
                    validacion.periodo,
                    directorio_ejecucion,
                    inspeccionar=inspeccionar_llm,
                    al_actualizar_progreso=publicar_progreso,
                    progreso_inicial=progreso_extraccion,
                    id_ejecucion=id_ejecucion,
                )
                decisiones_llm = analisis_llm.decisiones
            except Exception as exc:
                hallazgos.append(
                    _hallazgo(
                        "ANALISTA_LLM_NO_DISPONIBLE",
                        "warning",
                        (
                            "El analista curricular no estuvo disponible; se conserva "
                            "el resultado determinista."
                        ),
                        validacion.archivo,
                        f"{type(exc).__name__}: {str(exc)[:200]}",
                    )
                )
                _escribir_json(
                    reportes / "analisis_llm.json",
                    {
                        "estado": "FALLBACK_DETERMINISTA",
                        "error": f"{type(exc).__name__}: {str(exc)[:300]}",
                        "decisiones_aceptadas": 0,
                    },
                )
                publicar_progreso(
                    replace(
                        progreso_extraccion,
                        fase="error",
                        reporte_final="disponible",
                    ).con_evento(
                        "El análisis LLM no estuvo disponible; continúa la salida determinista."
                    )
                )
        resultado_catalogo = construir_salidas_curriculares(
            registros,
            validacion,
            directorio_ejecucion,
            catalogo_base,
            catalogo_carrera,
            decisiones_llm,
        )
        if analisis_llm is not None:
            _escribir_jsonl(reportes / "decisiones_llm.jsonl", analisis_llm.reportes)
            _escribir_json(
                reportes / "analisis_llm.json",
                {
                    "estado": "COMPLETADO",
                    "modelo_analista": analisis_llm.modelo_analista,
                    "modelo_inspector": analisis_llm.modelo_inspector,
                    "modelo_analista_residual": analisis_llm.modelo_analista_residual,
                    "modelo_inspector_residual": analisis_llm.modelo_inspector_residual,
                    "lotes": analisis_llm.lotes,
                    "decisiones_aceptadas": len(analisis_llm.decisiones),
                    "decisiones_escaladas": analisis_llm.decisiones_escaladas,
                    "decisiones_reportadas": len(analisis_llm.reportes),
                    "auditoria_contexto": analisis_llm.auditoria_contexto,
                },
            )
            if usar_llm:
                publicar_progreso(
                    replace(
                        progreso_extraccion,
                        fase="completado",
                        reporte_final="disponible",
                    ).con_evento("Reporte LLM disponible.")
                )
            report_outputs = tuple(
                resultado_catalogo.outputs
                + (
                    _output(
                        reportes / "decisiones_llm.jsonl",
                        "auditoria_llm",
                        len(analisis_llm.reportes),
                    ),
                    _output(reportes / "analisis_llm.json", "auditoria_llm", 1),
                )
            )
            resultado_catalogo = replace(resultado_catalogo, outputs=report_outputs)
    except Exception as exc:
        if usar_llm:
            publicar_progreso(
                replace(
                    progreso_extraccion,
                    fase="error",
                    reporte_final="error",
                ).con_evento(
                    "No se pudo completar el reporte LLM; se conservan los avances previos."
                )
            )
        hallazgo = _hallazgo(
            "CATALOGO_CURRICULAR_NO_DISPONIBLE",
            "error",
            "No se pudo construir el catálogo curricular con los catálogos base.",
            validacion.archivo,
            f"{type(exc).__name__}: {str(exc)[:200]}",
        )
        hallazgos.append(hallazgo)
        resultado_catalogo = ResultadoCatalogoCurricular(
            publicable=False,
            relaciones=0,
            competencias=0,
            habilidades=0,
            herramientas=0,
            outputs=(),
            hallazgos=(),
            cuarentena=(),
        )

    cuarentena.extend(resultado_catalogo.cuarentena)
    cuarentena_path = reportes / "cuarentena.jsonl"
    with cuarentena_path.open("w", encoding="utf-8", newline="\n") as salida_cuarentena:
        for fila in cuarentena:
            salida_cuarentena.write(
                json.dumps(fila, ensure_ascii=False, separators=(",", ":")) + "\n"
            )

    hallazgos_totales = tuple(hallazgos) + resultado_catalogo.hallazgos
    return ResultadoLimpiezaSilabos(
        registros=len(registros),
        outputs=resultado_catalogo.outputs,
        hallazgos=hallazgos_totales,
        publicable=resultado_catalogo.publicable,
        relaciones=resultado_catalogo.relaciones,
        competencias=resultado_catalogo.competencias,
        habilidades=resultado_catalogo.habilidades,
        herramientas=resultado_catalogo.herramientas,
    )


def _materializar(
    ruta_entrada: Path,
    directorio: Path,
    archivos: tuple[ArchivoSilabo, ...],
) -> dict[str, Path]:
    resultado: dict[str, Path] = {}
    if ruta_entrada.suffix.lower() != ".zip":
        nombre = archivos[0].nombre
        destino = directorio / Path(nombre).name
        shutil.copyfile(ruta_entrada, destino)
        resultado[nombre] = destino
        return resultado
    with zipfile.ZipFile(ruta_entrada) as paquete:
        for archivo in archivos:
            destino = directorio / archivo.nombre
            destino.parent.mkdir(parents=True, exist_ok=True)
            with (
                paquete.open(archivo.nombre.replace("\\", "/")) as origen,
                destino.open("wb") as salida,
            ):
                shutil.copyfileobj(origen, salida)
            resultado[archivo.nombre] = destino
    return resultado


def _primer_metadata(metadata: dict[str, str], claves: tuple[str, ...]) -> str:
    for clave in claves:
        if metadata.get(clave):
            return metadata[clave]
    return ""


def _nombre_desde_archivo(nombre: str) -> str:
    base = Path(nombre).stem.replace("_", " ").replace("-", " ")
    return re.sub(r"\s+", " ", base).strip().title()


def _ciclo_desde_ruta(nombre: str) -> str:
    coincidencia = re.search(r"ciclo[_ -]?(\d+)", nombre, re.IGNORECASE)
    return coincidencia.group(1) if coincidencia else ""


def _escribir_json(ruta: Path, contenido: dict[str, object]) -> None:
    ruta.write_text(
        json.dumps(contenido, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _escribir_jsonl(ruta: Path, filas: tuple[dict[str, object], ...]) -> None:
    with ruta.open("w", encoding="utf-8", newline="\n") as archivo:
        for fila in filas:
            archivo.write(json.dumps(fila, ensure_ascii=False, separators=(",", ":")))
            archivo.write("\n")


def _output(ruta: Path, tipo: str, registros: int) -> dict[str, object]:
    digest = hashlib.sha256()
    with ruta.open("rb") as archivo:
        for bloque in iter(lambda: archivo.read(1024 * 1024), b""):
            digest.update(bloque)
    ejecucion = next(
        (padre for padre in ruta.parents if padre.name.startswith("NOR_")),
        None,
    )
    return {
        "tipo": tipo,
        "archivo": str(ruta.relative_to(ejecucion)) if ejecucion else ruta.name,
        "registros": registros,
        "bytes": ruta.stat().st_size,
        "sha256": digest.hexdigest(),
    }
